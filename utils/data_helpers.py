from typing import Optional
import streamlit as st
import json
import csv
import io
from datetime import datetime, timedelta, date
from docx import Document
import pandas as pd
import plotly.express as px
from uuid import uuid4
import re
from collections import defaultdict
from config.sections import SECTION_METADATA

DEFAULT_COMMON_SECTIONS = set(SECTION_METADATA.keys())

def get_or_create_session_id():
    """Assigns a persistent unique ID per session"""
    if "session_id" not in st.session_state:
        st.session_state["session_id"] = str(uuid4())
    return st.session_state["session_id"]

def set_user_id(user_id):
    """Lets you manually tag a user (e.g., via login form or input"""
    st.session_state["user_id"] = user_id

def init_section(section_name):
    """Initialize a new section to group questions and answers."""
    if "input_data" not in st.session_state:
        st.session_state["input_data"] = {}
    if section_name not in st.session_state["input_data"]:
        st.session_state["input_data"][section_name] = []

def log_interaction(action_type, label, value, section_name):
    """Log a user interaction for troubleshooting or auditing."""
    if "interaction_log" not in st.session_state:
        st.session_state["interaction_log"] = []
    st.session_state["interaction_log"].append({
        "timestamp": datetime.now().isoformat(),
        "action": action_type,
        "question": label,
        "answer": value,
        "section": section_name
    })

def capture_input(
    label: str,
    input_fn,
    section: str,
    *args,
    validate_fn=None,
    preprocess_fn=None,
    required: bool = False,
    autofill: bool = True,
    metadata: dict = None,
    **kwargs
):
    """
    Captures user input and records structured metadata in session state.

    Returns:
        The final processed input value.
    """
    # Use explicit key or generate one
    default_key = f"{section}_{sanitize_label(label)}"
    unique_key = kwargs.get("key", default_key)
    kwargs["key"] = unique_key

    # Auto-fill value from get_answer()
    if autofill:
        default_value = get_answer(key=label, section=section)
        if default_value is not None and "value" not in kwargs:
            kwargs["value"] = default_value

        # Prevent .value for incompatible widgets
        if input_fn in [st.radio, st.selectbox, st.multiselect] and "value" in kwargs:
            kwargs.pop("value")

    # Render input widget
    value = input_fn(label, *args, **kwargs)

    # Optional transform
    if preprocess_fn:
        try:
            value = preprocess_fn(value)
        except Exception as e:
            st.error(f"❌ Error processing input: {e}")
            return None

    # Optional validation
    if validate_fn:
        try:
            if not validate_fn(value):
                st.warning(f"⚠️ Invalid value for '{label}'")
                return None
        except Exception as e:
            st.error(f"❌ Validation failed: {e}")
            return None

    init_section(section)

    entry = {
        "question": label,
        "answer": value,
        "key": unique_key,  # ✅ required for get_answer to match
        "timestamp": datetime.now().isoformat(),
        "input_type": getattr(input_fn, "__name__", str(input_fn)),
        "section": section,
        "session_id": st.session_state.get("session_id"),
        "required": required,
        "metadata": metadata or {},
    }

    st.session_state.setdefault("input_data", {}).setdefault(section, [])

    # ✅ Optional: overwrite previous entry for same key
    section_entries = st.session_state["input_data"][section]
    section_entries = [e for e in section_entries if e.get("key") != unique_key]
    section_entries.append(entry)
    st.session_state["input_data"][section] = section_entries

    log_interaction("input", label, value, section)
    autosave_input_data()

    return value

# Wrapper for task-based inputs
def register_task_input(
    label,
    input_fn,
    section: str,
    is_freq: bool = False,
    task_type: str = None,
    key: str = None,  # ✅ Explicit key
    *args,
    **kwargs
):
    """
    Registers a user input and stores it in session state for both tasks and non-tasks.
    Compatible with get_answer() by storing question/answer metadata in input_data.

    Args:
        label (str): The input label (used for question and task description).
        input_fn: Streamlit input function (e.g., st.text_area).
        section (str): The section name (flat string, e.g., "emergency_kit").
        is_freq (bool): Whether this task is frequency-based.
        task_type (str): Optional logical type (e.g., 'mail', 'trash').
        key (str): Optional key to ensure matching across inputs and get_answer().
        *args, **kwargs: Passed through to capture_input().
    """
    metadata = {
        "is_task": True,
        "task_label": label,
        "is_freq": is_freq,
        "area": "home",
        "section": section,
        "task_type": task_type,
    }
    kwargs["metadata"] = metadata

    if key:
        kwargs["key"] = key

    value = capture_input(label, input_fn, section=section, *args, **kwargs)

    if value not in (None, ""):
        if metadata["is_task"]:
            task_row = {
                "question": label,
                "answer": value,
                "category": section,
                "section": section,
                "area": metadata["area"],
                "task_type": metadata["task_type"],
                "is_freq": metadata["is_freq"],
                "key": key,  # ✅ Store key directly
            }
            st.session_state.setdefault("task_inputs", []).append(task_row)

        input_record = {
            "question": label,
            "answer": value,
            "section": section,
            "key": key,  # ✅ Store key here too
        }
        st.session_state.setdefault("input_data", {}).setdefault(section, []).append(input_record)

    return value

def sanitize(text):
    if not isinstance(text, str):
        print(f"[DEBUG] sanitize received invalid type: {text} ({type(text)})")
        return ""
    return (
        text.lower().strip()
        .replace(" ", "_")
        .replace("?", "")
        .replace(":", "")
        .replace("-", "_")
    )

def get_answer(
    *,
    key: str,
    section: str,
    nested_parent: str = None,
    nested_child: str = None,
    verbose: bool = False,
    common_sections: set = None
) -> str | None:
    """
    Retrieves the most recent answer for a given key in a section.
    Supports flat lookup, nested fallback, and optional verbose debugging.
    Nested_parent/nested_child is prioritized before task_inputs or input_data

    Args:
        section (str): Logical section name (e.g., "home", "emergency_kit").
        key (str): The question label or key to retrieve.
        nested_parent (str): Optional session_state parent (e.g., "input_data").
        nested_child (str): Optional nested dict under parent.
        verbose (bool): If True, shows debug output.
        common_sections (set): Optional override set of valid section names.

    Returns:
        str | None: Answer if found.
    """
    norm_key = sanitize(key)

    # Warn if arguments may be reversed
    common_set = common_sections or DEFAULT_COMMON_SECTIONS
    if section not in common_set and key in common_set:
        if verbose:
            st.warning(f"⚠️ `section='{section}'` may be reversed with `key='{key}'`")

    # Optional nested lookup (e.g., input_data["home"]["city"])
    if nested_parent and nested_child:
        nested_val = (
            st.session_state.get(nested_parent, {})
            .get(nested_child, {})
            .get(key)
        )
        if nested_val is not None:
            if verbose:
                st.write(f"✅ [Nested Fallback] Found in `{nested_parent} → {nested_child} → {key}`")
            return nested_val

    # 1️⃣ Look in st.session_state["task_inputs"]
    entries = st.session_state.get("task_inputs", [])
    for entry in entries:
        if entry.get("section") != section:
            continue
        label_raw = entry.get("question", "")
        key_raw = entry.get("key", "")

        if sanitize(label_raw) == norm_key or sanitize(key_raw) == norm_key:
            if verbose:
                st.write("🔍 Searching in `task_inputs`")
                st.write(f"✅ Match found — label: `{label_raw}`, key: `{key_raw}`")
            return entry.get("answer")

    # 2️⃣ Fallback to st.session_state["input_data"]
    input_data = st.session_state.get("input_data", {}).get(section, [])
    for entry in input_data:
        label_raw = entry.get("question", "")
        key_raw = entry.get("key", "")

        if sanitize(label_raw) == norm_key or sanitize(key_raw) == norm_key:
            if verbose:
                st.write("🔍 Searching in `input_data`")
                st.write(f"✅ Match found — label: `{label_raw}`, key: `{key_raw}`")
            return entry.get("answer")

    if verbose:
        st.warning(f"❌ No match found for key '{key}' in section '{section}'.")
    return None


def check_missing_utility_inputs(section="home"):
    """
    Checks for missing required fields for utility runbook generation.
    Returns a list of missing field labels.
    """
    missing = []

    if not get_answer(key="City", section=section):
        missing.append("City")
    if not get_answer(key="ZIP Code", section=section):
        missing.append("ZIP Code")
    if not get_answer(key="Internet Provider", section=section):
        missing.append("Internet Provider")

    return missing

def flatten_answers_to_dict(questions, answers):
    """
    Converts a flat list of questions and answers into a dictionary,
    filtering out empty answers.
    """
    return {
        question: answer
        for question, answer in zip(questions, answers)
        if str(answer).strip()
    }

def preview_input_data():
    """Display a summary of all collected input data."""
    if "input_data" not in st.session_state or not st.session_state["input_data"]:
        st.warning("No input data collected yet.")
        return
    st.markdown("### 📝 Input Summary")
    for section, entries in st.session_state["input_data"].items():
        st.markdown(f"#### 📁 {section}")
        for item in entries:
            st.markdown(f"- **{item['question']}**: {item['answer']} _(at {item['timestamp']})_")

def daterange(start, end):
    for n in range((end - start).days + 1):
        yield start + timedelta(n)

def get_filtered_dates(start_date, end_date, refinement):
    if refinement == "Weekdays Only":
        return [d for d in daterange(start_date, end_date) if d.weekday() < 5]
    elif refinement == "Weekend Only":
        return [d for d in daterange(start_date, end_date) if d.weekday() >= 5]
    else:  # All Days
        return list(daterange(start_date, end_date))
    
def select_runbook_date_range(section="runbook_date_range"):
    st.write("📅 First Select Date(s) or Timeframe for Your Runbook.")

    options = ["Pick Dates", "General"]
    choice = register_task_input(
        label="Runbook Timeframe Option",
        input_fn=st.radio,
        section=section,
        options=options,
        index=0,
        key="runbook_timeframe_option"
    )

    start_date, end_date = None, None
    valid_dates = []
    today = datetime.now().date()

    if choice == "Pick Dates":
        start_date = register_task_input(
            label="Start Date",
            input_fn=st.date_input,
            section=section,
            value=today,
            key="start_date_input"
        )
        end_date = register_task_input(
            label="End Date",
            input_fn=st.date_input,
            section=section,
            value=today + timedelta(days=7),
            key="end_date_input"
        )

        if start_date >= end_date:
            st.error("⚠️ Start date must be before end date.")
            return None, None, None, []

        days_apart = (end_date - start_date).days
        if days_apart > 31:
            st.error(f"⚠️ The selected period is too long ({days_apart} days). Limit: 31 days.")
            return None, None, None, []

        refinement = register_task_input(
            label="Date Range Refinement",
            input_fn=st.radio,
            section=section,
            options=["All Days", "Weekdays Only", "Weekend Only"],
            index=0,
            horizontal=True,
            key="refinement_option"
        )

        display_choice = f"Pick Dates ({refinement})"
        st.info(f"📅 Using dates from **{start_date}** to **{end_date}** ({refinement})")
        valid_dates = get_filtered_dates(start_date, end_date, refinement)

    elif choice == "General":
        start_date = today
        end_date = today + timedelta(days=7)
        display_choice = "General (Next 7 Days)"
        st.info(f"📅 1-week schedule starting {start_date}")
        valid_dates = get_filtered_dates(start_date, end_date, "All Days")

    else:
        st.warning("⚠️ Invalid choice selected.")
        return None, None, None, []

    return display_choice, start_date, end_date, valid_dates

def preview_interaction_log():
    """Display a log of all user interactions."""
    if "interaction_log" not in st.session_state or not st.session_state["interaction_log"]:
        st.info("No interactions logged yet.")
        return
    st.markdown("### 🧾 Interaction History")
    for log in st.session_state["interaction_log"]:
        st.markdown(f"- [{log['timestamp']}] {log['action'].capitalize()} — **{log['question']}** → `{log['answer']}` _(Section: {log['section']})_")

def custom_serializer(obj):
    # Handle datetime and date objects
    if isinstance(obj, (datetime, date)):
        return obj.isoformat()  # ✅ Convert to ISO string like "2025-05-24"
    
    # If it's some other unsupported type, raise an error
    raise TypeError(f"Type {type(obj)} not serializable")

def autosave_input_data():
    """Automatically save data to session file (local or cloud placeholder)."""
    if "input_data" in st.session_state:
        st.session_state["autosaved_json"] = json.dumps(
        st.session_state["input_data"],
        indent=2,
        default=custom_serializer  # ✅ use your custom handler
        )
        # Placeholder: save_to_cloud(st.session_state["autosaved_json"])

def export_input_data_as_json(file_name="input_data.json"):
    """Export the collected input data as JSON."""
    if "input_data" in st.session_state:
        json_data = json.dumps(
            st.session_state["input_data"], 
            indent=2, 
            default=custom_serializer # ✅ Handles date/datetime
            )
        st.download_button(
            label="📥 Download as JSON",
            data=json_data,
            file_name=file_name,
            mime="application/json"
        )

def is_iso_date(string: str) -> bool:
    """Check if a string matches an ISO 8601 date format (YYYY-MM-DD)."""
    try:
        datetime.strptime(string, "%Y-%m-%d")
        return True
    except ValueError:
        return False

def restore_dates(obj):
    """Recursively convert ISO date strings into datetime.date objects."""
    if isinstance(obj, dict):
        return {k: restore_dates(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [restore_dates(item) for item in obj]
    elif isinstance(obj, str) and is_iso_date(obj):
        return datetime.strptime(obj, "%Y-%m-%d").date()
    return obj

def import_input_data_from_json(json_str: str):
    """
    Load JSON string, convert ISO-formatted dates back into date objects,
    and save into session state.
    """
    try:
        parsed = json.loads(json_str)
        restored = restore_dates(parsed)
        st.session_state["input_data"] = restored
        st.success("✅ Input data successfully imported and restored.")
    except Exception as e:
        st.error(f"❌ Failed to import input data: {e}")

# Streamlit code usage to re-import JSON data and automatically convert ISO date strings back into datetime.date objects.
#uploaded_file = st.file_uploader("Upload previously exported input JSON", type=["json"])
#if uploaded_file:
#    json_str = uploaded_file.read().decode("utf-8")
#    if st.button("📤 Import JSON Data"):
#       import_input_data_from_json(json_str)

def convert_to_csv(data_list):
    if not data_list:
        return ""

    # Collect all possible keys
    fieldnames = sorted({key for row in data_list for key in row})

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()

    for row in data_list:
        row_copy = row.copy()
        for k, v in row_copy.items():
            if isinstance(v, dict):
                row_copy[k] = json.dumps(v)
        writer.writerow(row_copy)

    return output.getvalue()

def export_input_data_as_csv(file_name="input_data.csv"):
    """Wraps convert_to_csv() to export all input_data with section info."""
    full_list = []
    for section, entries in st.session_state.get("input_data", {}).items():
        for item in entries:
            item_copy = item.copy()
            item_copy["section"] = section
            full_list.append(item_copy)

    csv_data = convert_to_csv(full_list)

    st.download_button(
        label="📄 Download Your Data",
        data=csv_data,
        file_name=file_name,
        mime="text/csv"
    )

def export_input_data_as_docx(file_name="input_data.docx"):
    """Export the collected input data as DOCX."""
    if "input_data" in st.session_state:
        doc = Document()
        doc.add_heading("Collected Input Summary", level=1)
        for section, entries in st.session_state["input_data"].items():
            doc.add_heading(section, level=2)
            for item in entries:
                doc.add_paragraph(f"{item['question']}: {item['answer']} ({item['timestamp']})")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📝 Download as DOCX",
            data=buffer,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def render_lock_toggle(section: str, session_key: Optional[str] = None, label: Optional[str] = None):
    """
    Renders a simple toggle switch to lock/unlock inputs and updates session state.

    Args:
        section (str): The logical section name (e.g., "mail", "trash_handling").
        session_key (Optional[str]): Optional override for session state key.
        label (Optional[str]): Optional display label for the section.
    """
    key = session_key or f"{section}_locked"
    display_label = label or section.replace("_", " ").title()

    if key not in st.session_state:
        st.session_state[key] = False  # default unlocked

    is_locked = st.toggle(f"🔒 Lock {display_label}", value=st.session_state[key])
    st.session_state[key] = is_locked

    if is_locked:
        st.success(f"✅ {display_label} is saved and locked. Unlock to edit.")
    else:
        st.info(f"📝 You can now edit your {display_label.lower()}. Lock to save when finished.")

def interaction_dashboard():
    if "interaction_log" not in st.session_state or not st.session_state["interaction_log"]:
        st.info("No interaction data to visualize.")
        return

    df = pd.DataFrame(st.session_state["interaction_log"])
    st.markdown("### 📊 Interaction Dashboard")

    section_filter = st.selectbox("Filter by section", options=["All"] + sorted(df["section"].unique().tolist()))
    if section_filter != "All":
        df = df[df["section"] == section_filter]

    st.dataframe(df)

    st.markdown("#### 🔄 Inputs per Section")
    fig = px.histogram(df, x="section", color="action", barmode="group")
    st.plotly_chart(fig, use_container_width=True)

    st.markdown("#### ⏱ Interaction Timeline")
    df['timestamp'] = pd.to_datetime(df['timestamp'])
    timeline = px.scatter(df, x="timestamp", y="section", color="user_id", hover_data=["question", "answer"])
    st.plotly_chart(timeline, use_container_width=True)

def log_provider_result(label, value, section="Utility Providers"):
    """
    Logs a single provider result into input_data, with basic validation.
    - Skips empty, 'not found', or 'n/a' values
    - Raises an error for invalid labels
    """
    if not label or not isinstance(label, str):
        raise ValueError("Provider label must be a non-empty string.")

    if not isinstance(value, str) or value.strip().lower() in ["", "not found", "n/a"]:
        return  # Skip logging invalid or placeholder values

    if "input_data" not in st.session_state:
        st.session_state["input_data"] = {}
    if section not in st.session_state["input_data"]:
        st.session_state["input_data"][section] = []

    st.session_state["input_data"][section].append({
        "question": f"{label} Provider",
        "answer": value,
        "timestamp": datetime.now().isoformat()
    })

    if "interaction_log" not in st.session_state:
        st.session_state["interaction_log"] = []
    session_id = st.session_state.get("session_id", "anonymous")
    user_id = st.session_state.get("user_id", "anonymous")

    st.session_state["interaction_log"].append({
        "timestamp": datetime.now().isoformat(),
        "user_id": user_id,
        "session_id": session_id,
        "action": "autolog",
        "question": f"{label} Provider",
        "answer": value,
        "section": section
    })

def extract_providers_from_text(content: str) -> dict:
    """
    Pure function that extracts provider names from a text string.
    Returns a dict with electricity, natural_gas, and water keys.
    """
    def extract(label):
        match = re.search(rf"{label} Provider:\s*(.+)", content, re.IGNORECASE)
        return match.group(1).strip() if match else "Not found"

    return {
        "electricity": extract("Electricity"),
        "natural_gas": extract("Natural Gas"),
        "water": extract("Water")
    }

def extract_and_log_providers(content: str) -> dict:
    """
    Wrapper that extracts provider names and logs them to session state.
    """
    providers = extract_providers_from_text(content)

    # Log to Streamlit or any side-effect mechanism
    log_provider_result("Electricity", providers["electricity"])
    log_provider_result("Natural Gas", providers["natural_gas"])
    log_provider_result("Water", providers["water"])

    # Store in session state for reuse
    st.session_state["electricity_provider"] = providers["electricity"]
    st.session_state["natural_gas_provider"] = providers["natural_gas"]
    st.session_state["water_provider"] = providers["water"]

    return providers

def generate_category_keywords_from_labels(input_data):
    """Auto-generate category keywords based on recurring words in labels."""
    keyword_freq = defaultdict(int)

    for section, entries in input_data.items():
        for entry in entries:
            label = entry.get("question", "")
            words = re.findall(r"\b\w{4,}\b", label.lower())  # Filter short/noisy words
            for word in words:
                keyword_freq[word] += 1

    # Filter common task-related words only (tunable)
    candidate_keywords = {k: v for k, v in keyword_freq.items() if v >= 2}  # show those used more than once
    sorted_keywords = sorted(candidate_keywords.items(), key=lambda x: -x[1])

    return sorted_keywords

def sanitize_label(label: str) -> str:
    return label.lower().strip().replace(" ", "_").replace("?", "").replace(":", "")

def section_has_valid_input(section: str, min_entries: int = 1) -> bool:
    """
    Check if a given section has at least `min_entries` valid task inputs in session state.

    Args:
        section (str): The section name to filter by (e.g. "mail_trash").
        min_entries (int): Minimum number of entries required to consider the section valid.

    Returns:
        bool: True if section has enough inputs, False otherwise.
    """
    inputs = st.session_state.get("task_inputs", [])
    valid = [item for item in inputs if item.get("section") == section and item.get("answer") not in ("", None)]
    return len(valid) >= min_entries




# Placeholder for future enhancement: Cloud saving logic
# def save_to_cloud(json_data):
#     """Save the JSON data to a cloud storage provider."""
#     pass
