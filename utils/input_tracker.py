import streamlit as st
import json
import csv
import io
from datetime import datetime, timedelta, date
from docx import Document
import pandas as pd
import plotly.express as px
from uuid import uuid4
import re
from collections import defaultdict

def get_or_create_session_id():
    """Assigns a persistent unique ID per session"""
    if "session_id" not in st.session_state:
        st.session_state["session_id"] = str(uuid4())
    return st.session_state["session_id"]

def set_user_id(user_id):
    """Lets you manually tag a user (e.g., via login form or input"""
    st.session_state["user_id"] = user_id

def init_section(section_name):
    """Initialize a new section to group questions and answers."""
    if "input_data" not in st.session_state:
        st.session_state["input_data"] = {}
    if section_name not in st.session_state["input_data"]:
        st.session_state["input_data"][section_name] = []

def log_interaction(action_type, label, value, section_name):
    """Log a user interaction for troubleshooting or auditing."""
    if "interaction_log" not in st.session_state:
        st.session_state["interaction_log"] = []
    st.session_state["interaction_log"].append({
        "timestamp": datetime.now().isoformat(),
        "action": action_type,
        "question": label,
        "answer": value,
        "section": section_name
    })

def capture_input(
    label,
    input_fn,
    section=None,
    *args,
    validate_fn=None,
    preprocess_fn=None,
    required=False,
    autofill=True,
    **kwargs
):
    if section is None:
        section = st.session_state.get("section", "default")

    unique_key = kwargs.get("key", f"{section}_{label.replace(' ', '_').lower()}")
    kwargs["key"] = unique_key

    # Auto-fill value using get_answer()
    default_value = None
    if autofill:
        default_value = get_answer(label, section)
        if default_value is not None and "value" not in kwargs:
            kwargs["value"] = default_value

        # 🛡️ Defensive cleanup
        if input_fn in [st.radio, st.selectbox] and "value" in kwargs:
            kwargs.pop("value")
        if input_fn == st.multiselect and "value" in kwargs:
            kwargs.pop("value")


    value = input_fn(label, *args, **kwargs)

    if preprocess_fn:
        try:
            value = preprocess_fn(value)
        except Exception as e:
            st.error(f"❌ Error processing input: {e}")
            return None

    if validate_fn:
        try:
            is_valid = validate_fn(value)
            if not is_valid:
                st.warning(f"⚠️ Invalid value for '{label}'")
                return None
        except Exception as e:
            st.error(f"❌ Validation failed: {e}")
            return None

    init_section(section)

    entry = {
        "question": label,
        "answer": value,
        "timestamp": datetime.now().isoformat(),
        "input_type": getattr(input_fn, "__name__", str(input_fn)),
        "section": section,
        "session_id": st.session_state.get("session_id"),
        "required": required,
    }

    if "input_data" not in st.session_state:
        st.session_state["input_data"] = {}
    st.session_state["input_data"].setdefault(section, []).append(entry)

    log_interaction("input", label, value, section)
    autosave_input_data()

    return value


def get_answer(question_label: str, section: str = None, *, nested_parent: str = None, nested_child: str = None):
    """
    Retrieves the most recent answer matching the question_label.
    
    Lookup Order:
    1. If nested_parent and nested_child are given, check nested structure.
    2. Otherwise, check flat input_data log scoped by section.

    Args:
        question_label (str): The exact label used during capture_input.
        section (str): The flat section name in input_data (optional).
        nested_parent (str): Top-level nested dict key (e.g. "quality_services").
        nested_child (str): Second-level nested dict key (e.g. "Cleaning").
    
    Returns:
        str | None: The matched answer, or None if not found.
    """
    # ✅ Check nested dict structure
    if nested_parent and nested_child:
        nested_value = (
            st.session_state.get(nested_parent, {})
            .get(nested_child, {})
            .get(question_label)
        )
        if nested_value is not None:
            return nested_value

    # ✅ Fallback to flat legacy input_data
    data = st.session_state.get("input_data", {})
    sections = [section] if section else data.keys()

    latest = None
    for sec in sections:
        for item in data.get(sec, []):
            if item["question"] == question_label:
                latest = item["answer"]  # override to keep the most recent
    return latest

def flatten_answers_to_dict(questions, answers):
    """
    Converts a flat list of questions and answers into a dictionary,
    filtering out empty answers.
    """
    return {
        question: answer
        for question, answer in zip(questions, answers)
        if str(answer).strip()
    }

def autolog_location_inputs(city, zip_code):
    """
    Optionally logs city and zip code inputs to 'Home Basics' section if not already logged.
    """
    now = datetime.now().isoformat()

    # Ensure input_data exists
    if "input_data" not in st.session_state:
        st.session_state["input_data"] = {}

    # Ensure the 'Home Basics' section exists
    if "Home Basics" not in st.session_state["input_data"]:
        st.session_state["input_data"]["Home Basics"] = []

    def log_if_missing(label, value):
        if value and not any(entry["question"] == label for entry in st.session_state["input_data"]["Home Basics"]):
            # Add to input_data
            st.session_state["input_data"]["Home Basics"].append({
                "question": label,
                "answer": value,
                "timestamp": now
            })

            # Also add to interaction_log if enabled
            st.session_state.setdefault("interaction_log", []).append({
                "timestamp": now,
                "user_id": st.session_state.get("user_id", "anonymous"),
                "session_id": st.session_state.get("session_id", "anonymous"),
                "action": "autolog",
                "question": label,
                "answer": value,
                "section": "Home Basics"
            })

    log_if_missing("City", city)
    log_if_missing("ZIP Code", zip_code)

def preview_input_data():
    """Display a summary of all collected input data."""
    if "input_data" not in st.session_state or not st.session_state["input_data"]:
        st.warning("No input data collected yet.")
        return
    st.markdown("### 📝 Input Summary")
    for section, entries in st.session_state["input_data"].items():
        st.markdown(f"#### 📁 {section}")
        for item in entries:
            st.markdown(f"- **{item['question']}**: {item['answer']} _(at {item['timestamp']})_")

def daterange(start, end):
    for n in range((end - start).days + 1):
        yield start + timedelta(n)

def get_filtered_dates(start_date, end_date, refinement):
    if refinement == "Weekdays Only":
        return [d for d in daterange(start_date, end_date) if d.weekday() < 5]
    elif refinement == "Weekend Only":
        return [d for d in daterange(start_date, end_date) if d.weekday() >= 5]
    else:  # All Days
        return list(daterange(start_date, end_date))
    
def select_runbook_date_range():
    section = "Runbook Date Range"
    #st.subheader("📅 Choose Date(s) or Timeframe")
    st.write("📅 First Select Date(s) or Timeframe for Your Runbook.")

    options = ["Pick Dates", "General"]
    choice = capture_input(
        "Runbook Timeframe Option",
        st.radio,
        section,
        options=options,
        index=0,
        key="runbook_timeframe_option"
    )

    start_date, end_date = None, None
    valid_dates = []
    today = datetime.now().date()

    if choice == "Pick Dates":
        start_date = capture_input(
            "Start Date",
            st.date_input,
            section,
            value=today,
            key="start_date_input"
        )
        end_date = capture_input(
            "End Date",
            st.date_input,
            section,
            value=today + timedelta(days=7),
            key="end_date_input"
        )

        if start_date >= end_date:
            st.error("⚠️ Start date must be before end date.")
            return None, None, None, []

        if (end_date - start_date).days > 31:
            st.error("⚠️ The selected period must be no longer than 1 month.")
            return None, None, None, []

        refinement = capture_input(
            "Date Range Refinement",
            st.radio,
            section,
            options=["All Days", "Weekdays Only", "Weekend Only"],
            index=0,
            horizontal=True,
            key="refinement_option"
        )

        st.info(f"📅 Using dates from **{start_date}** to **{end_date}** ({refinement})")
        valid_dates = get_filtered_dates(start_date, end_date, refinement)
        choice = f"Pick Dates ({refinement})"

    elif choice == "General":
        start_date = today
        end_date = today + timedelta(days=7)
        st.info(f"📅 1-week schedule starting {start_date}")
        valid_dates = get_filtered_dates(start_date, end_date, "All Days")

    else:
        st.warning("⚠️ Invalid choice selected.")
        return None, None, None, []

    return choice, start_date, end_date, valid_dates

def preview_interaction_log():
    """Display a log of all user interactions."""
    if "interaction_log" not in st.session_state or not st.session_state["interaction_log"]:
        st.info("No interactions logged yet.")
        return
    st.markdown("### 🧾 Interaction History")
    for log in st.session_state["interaction_log"]:
        st.markdown(f"- [{log['timestamp']}] {log['action'].capitalize()} — **{log['question']}** → `{log['answer']}` _(Section: {log['section']})_")

def custom_serializer(obj):
    # Handle datetime and date objects
    if isinstance(obj, (datetime, date)):
        return obj.isoformat()  # ✅ Convert to ISO string like "2025-05-24"
    
    # If it's some other unsupported type, raise an error
    raise TypeError(f"Type {type(obj)} not serializable")

def autosave_input_data():
    """Automatically save data to session file (local or cloud placeholder)."""
    if "input_data" in st.session_state:
        st.session_state["autosaved_json"] = json.dumps(
        st.session_state["input_data"],
        indent=2,
        default=custom_serializer  # ✅ use your custom handler
        )
        # Placeholder: save_to_cloud(st.session_state["autosaved_json"])

def export_input_data_as_json(file_name="input_data.json"):
    """Export the collected input data as JSON."""
    if "input_data" in st.session_state:
        json_data = json.dumps(
            st.session_state["input_data"], 
            indent=2, 
            default=custom_serializer # ✅ Handles date/datetime
            )
        st.download_button(
            label="📥 Download as JSON",
            data=json_data,
            file_name=file_name,
            mime="application/json"
        )

def is_iso_date(string: str) -> bool:
    """Check if a string matches an ISO 8601 date format (YYYY-MM-DD)."""
    try:
        datetime.strptime(string, "%Y-%m-%d")
        return True
    except ValueError:
        return False

def restore_dates(obj):
    """Recursively convert ISO date strings into datetime.date objects."""
    if isinstance(obj, dict):
        return {k: restore_dates(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [restore_dates(item) for item in obj]
    elif isinstance(obj, str) and is_iso_date(obj):
        return datetime.strptime(obj, "%Y-%m-%d").date()
    return obj

def import_input_data_from_json(json_str: str):
    """
    Load JSON string, convert ISO-formatted dates back into date objects,
    and save into session state.
    """
    try:
        parsed = json.loads(json_str)
        restored = restore_dates(parsed)
        st.session_state["input_data"] = restored
        st.success("✅ Input data successfully imported and restored.")
    except Exception as e:
        st.error(f"❌ Failed to import input data: {e}")

# Streamlit code usage to re-import JSON data and automatically convert ISO date strings back into datetime.date objects.
#uploaded_file = st.file_uploader("Upload previously exported input JSON", type=["json"])
#if uploaded_file:
#    json_str = uploaded_file.read().decode("utf-8")
#    if st.button("📤 Import JSON Data"):
#       import_input_data_from_json(json_str)

def convert_to_csv(data_list):
    if not data_list:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=data_list[0].keys())
    writer.writeheader()
    writer.writerows(data_list)
    return output.getvalue()

def export_input_data_as_csv(file_name="input_data.csv"):
    """Wraps convert_to_csv() to export all input_data with section info."""
    full_list = []
    for section, entries in st.session_state.get("input_data", {}).items():
        for item in entries:
            item_copy = item.copy()
            item_copy["section"] = section
            full_list.append(item_copy)

    csv_data = convert_to_csv(full_list)

    st.download_button(
        label="📄 Download Your Data",
        data=csv_data,
        file_name=file_name,
        mime="text/csv"
    )

def export_input_data_as_docx(file_name="input_data.docx"):
    """Export the collected input data as DOCX."""
    if "input_data" in st.session_state:
        doc = Document()
        doc.add_heading("Collected Input Summary", level=1)
        for section, entries in st.session_state["input_data"].items():
            doc.add_heading(section, level=2)
            for item in entries:
                doc.add_paragraph(f"{item['question']}: {item['answer']} ({item['timestamp']})")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📝 Download as DOCX",
            data=buffer,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def render_lock_toggle(session_key: str, label: str = "Section"):
    """
    Renders a simple toggle switch to lock/unlock inputs and updates session state.

    Args:
        session_key (str): The key in st.session_state to track lock state.
        label (str): The label to describe the section being controlled.
    """
    if session_key not in st.session_state:
        st.session_state[session_key] = False  # default unlocked

    # Display the toggle
    is_locked = st.toggle(f"🔒 Lock {label}", value=st.session_state[session_key])
    st.session_state[session_key] = is_locked

    # Show contextual feedback
    if is_locked:
        st.success(f"✅ {label} is saved and locked. Unlock to edit.")
    else:
        st.info(f"📝 You can now edit your {label.lower()}. Lock to save when finished.")

def interaction_dashboard():
    if "interaction_log" not in st.session_state or not st.session_state["interaction_log"]:
        st.info("No interaction data to visualize.")
        return

    df = pd.DataFrame(st.session_state["interaction_log"])
    st.markdown("### 📊 Interaction Dashboard")

    section_filter = st.selectbox("Filter by section", options=["All"] + sorted(df["section"].unique().tolist()))
    if section_filter != "All":
        df = df[df["section"] == section_filter]

    st.dataframe(df)

    st.markdown("#### 🔄 Inputs per Section")
    fig = px.histogram(df, x="section", color="action", barmode="group")
    st.plotly_chart(fig, use_container_width=True)

    st.markdown("#### ⏱ Interaction Timeline")
    df['timestamp'] = pd.to_datetime(df['timestamp'])
    timeline = px.scatter(df, x="timestamp", y="section", color="user_id", hover_data=["question", "answer"])
    st.plotly_chart(timeline, use_container_width=True)

def log_provider_result(label, value, section="Utility Providers"):
    """
    Logs a single provider result into input_data, with basic validation.
    - Skips empty, 'not found', or 'n/a' values
    - Raises an error for invalid labels
    """
    if not label or not isinstance(label, str):
        raise ValueError("Provider label must be a non-empty string.")

    if not isinstance(value, str) or value.strip().lower() in ["", "not found", "n/a"]:
        return  # Skip logging invalid or placeholder values

    if "input_data" not in st.session_state:
        st.session_state["input_data"] = {}
    if section not in st.session_state["input_data"]:
        st.session_state["input_data"][section] = []

    st.session_state["input_data"][section].append({
        "question": f"{label} Provider",
        "answer": value,
        "timestamp": datetime.now().isoformat()
    })

    if "interaction_log" not in st.session_state:
        st.session_state["interaction_log"] = []
    session_id = st.session_state.get("session_id", "anonymous")
    user_id = st.session_state.get("user_id", "anonymous")

    st.session_state["interaction_log"].append({
        "timestamp": datetime.now().isoformat(),
        "user_id": user_id,
        "session_id": session_id,
        "action": "autolog",
        "question": f"{label} Provider",
        "answer": value,
        "section": section
    })

def extract_providers_from_text(content: str) -> dict:
    """
    Pure function that extracts provider names from a text string.
    Returns a dict with electricity, natural_gas, and water keys.
    """
    def extract(label):
        match = re.search(rf"{label} Provider:\s*(.+)", content, re.IGNORECASE)
        return match.group(1).strip() if match else "Not found"

    return {
        "electricity": extract("Electricity"),
        "natural_gas": extract("Natural Gas"),
        "water": extract("Water")
    }

def extract_and_log_providers(content: str) -> dict:
    """
    Wrapper that extracts provider names and logs them to session state.
    """
    providers = extract_providers_from_text(content)

    # Log to Streamlit or any side-effect mechanism
    log_provider_result("Electricity", providers["electricity"])
    log_provider_result("Natural Gas", providers["natural_gas"])
    log_provider_result("Water", providers["water"])

    # Store in session state for reuse
    st.session_state["electricity_provider"] = providers["electricity"]
    st.session_state["natural_gas_provider"] = providers["natural_gas"]
    st.session_state["water_provider"] = providers["water"]

    return providers

def check_missing_utility_inputs():
    """
    Checks for missing required fields for utility runbook generation.
    Returns a list of missing field labels.
    """

    missing = []

    # User inputs (Home Basics)
    if not get_answer("City", "Home Basics"):
        missing.append("City")
    if not get_answer("ZIP Code", "Home Basics"):
        missing.append("ZIP Code")
    if not get_answer("Internet Provider", "Home Basics"):
        missing.append("Internet Provider")

    # AI/Corrected utility providers (Utility Providers)
    if not get_answer("Electricity Provider", "Utility Providers"):
        missing.append("Electricity Provider")
    if not get_answer("Natural Gas Provider", "Utility Providers"):
        missing.append("Natural Gas Provider")
    if not get_answer("Water Provider", "Utility Providers"):
        missing.append("Water Provider")

    return missing

def generate_category_keywords_from_labels(input_data):
    """Auto-generate category keywords based on recurring words in labels."""
    keyword_freq = defaultdict(int)

    for section, entries in input_data.items():
        for entry in entries:
            label = entry.get("question", "")
            words = re.findall(r"\b\w{4,}\b", label.lower())  # Filter short/noisy words
            for word in words:
                keyword_freq[word] += 1

    # Filter common task-related words only (tunable)
    candidate_keywords = {k: v for k, v in keyword_freq.items() if v >= 2}  # show those used more than once
    sorted_keywords = sorted(candidate_keywords.items(), key=lambda x: -x[1])

    return sorted_keywords



# Placeholder for future enhancement: Cloud saving logic
# def save_to_cloud(json_data):
#     """Save the JSON data to a cloud storage provider."""
#     pass
