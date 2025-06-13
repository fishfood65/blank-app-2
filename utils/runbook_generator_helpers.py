import streamlit as st
import csv
from io import BytesIO, StringIO
import io
from datetime import datetime, timedelta
from collections import defaultdict
import json
import os
import pandas as pd
import re
from typing import Callable, List, Literal, Tuple, Optional, Union
import markdown as md
import tempfile
from .common_helpers import get_schedule_placeholder_mapping
from .prompt_block_utils import generate_all_prompt_blocks
from .task_schedule_utils_updated import export_schedule_to_markdown
import markdown
from config.constants import TASK_TYPE_EMOJI, SCHEDULE_HEADING_MAP, PRIORITY_ORDER
import pandas as pd
from datetime import datetime
import tiktoken  # For OpenAI-compatible token counting
import time
import traceback

def display_user_friendly_schedule_table(
    df: pd.DataFrame,
    label_col: str = "clean_task",
    show_heading: bool = True,
    heading_text: str = "📆 Scheduled Tasks",
    show_legend: bool = True,
    enable_task_filter: bool = True
):
    """
    Displays a user-friendly task schedule table:
    - Sorts by date and priority
    - Adds emojis and task labels
    - Gracefully handles missing columns
    - Includes collapsible emoji legend and task type filter
    """

    if df is None or df.empty:
        st.warning("⚠️ No scheduled tasks to display.")
        return

    df = df.copy()

    # ✅ Optional task type filter (applied early)
    if enable_task_filter and "task_type" in df.columns:
        with st.expander("🎛️ Filter by Task Type", expanded=False):
            unique_types = sorted(df["task_type"].dropna().unique())
            selected_types = st.multiselect("Show only these task types:", unique_types, default=unique_types)
            df = df[df["task_type"].isin(selected_types)]

    # 🏷️ Add emoji and readable label
    if "task_type" in df.columns:
        df["emoji"] = df["task_type"].map(TASK_TYPE_EMOJI).fillna("❓")
        df["🏷️ Task Type"] = df["emoji"] + " " + df["task_type"]
    else:
        df["🏷️ Task Type"] = "❓ Unknown"

    # 🧮 Add priority
    if "task_type" in df.columns:
        priority_lookup = {task: i for i, task in enumerate(PRIORITY_ORDER, 1)}
        df["priority"] = df["task_type"].map(priority_lookup).fillna(len(PRIORITY_ORDER) + 1).astype(int)
    else:
        df["priority"] = len(PRIORITY_ORDER) + 1

    # 🧮 Sort by Date and priority
    if "Date" in df.columns:
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        df = df.sort_values(by=["Date", "priority"])

    # 📋 Column label mapping
    display_cols = {}

    if "Date" in df.columns:
        display_cols["Date"] = "📅 Date"
    if "Day" in df.columns:
        display_cols["Day"] = "🗓️ Day"

    display_cols["🏷️ Task Type"] = "🏷️ Task Type"

    if label_col in df.columns:
        display_cols[label_col] = "✅ Task"
    else:
        df[label_col] = "⚠️ Missing"
        display_cols[label_col] = "✅ Task"

    user_friendly_df = df[list(display_cols.keys())].rename(columns=display_cols)

    # 🧾 Final display
    if show_heading:
        st.subheader(heading_text)

    st.dataframe(user_friendly_df, use_container_width=True)

def add_table_from_schedule(doc, schedule_df, section: str, include_priority: bool = True, include_heading: bool = False, heading_text: Optional[str] = None):
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if schedule_df.empty:
        return

    schedule_df = schedule_df.copy()
    schedule_df["Date"] = pd.to_datetime(schedule_df["Date"], errors="coerce")
    schedule_df["Priority"] = schedule_df["task_type"].apply(
        lambda t: PRIORITY_ORDER.index(t) if t in PRIORITY_ORDER else len(PRIORITY_ORDER)
    )
    schedule_df["Task Type"] = schedule_df["task_type"].apply(
        lambda t: f"{TASK_TYPE_EMOJI.get(t, '')} {t}" if t else "❓"
    )

    schedule_df = schedule_df.sort_values(by=["Date", "Priority", "clean_task"])

    # 🔹 Add main schedule heading (level 2) once before any date blocks
    if include_heading and heading_text:
        doc.add_heading(heading_text, level=2)
        doc.add_paragraph("")  # space after heading

    # 🔹 Add each date block under heading level 3
    for date, group in schedule_df.groupby("Date"):
        day_str = date.strftime("%A, %Y-%m-%d")
        doc.add_heading(f"📅 {day_str}", level=3)

        table = doc.add_table(rows=1, cols=3 if include_priority else 2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = "📝 Task"
        if include_priority:
            hdr_cells[1].text = "🔢 Priority"
            hdr_cells[2].text = "🔖 Type"
        else:
            hdr_cells[1].text = "🔖 Type"

        for _, row in group.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row.get("clean_task", ""))
            if include_priority:
                cells[1].text = str(row.get("Priority", ""))
                cells[2].text = str(row.get("Task Type", ""))
            else:
                cells[1].text = str(row.get("Task Type", ""))

        doc.add_paragraph("")

def add_table_from_schedule_to_markdown(schedule_df: pd.DataFrame, section: str, include_priority: bool = True, include_heading: bool = False, heading_text: Optional[str] = None) -> str:
    if schedule_df.empty:
        return "_No schedule data available._"

    schedule_df = schedule_df.copy()
    schedule_df["Date"] = pd.to_datetime(schedule_df["Date"], errors="coerce")
    schedule_df["Day"] = schedule_df["Date"].dt.strftime("%A")
    schedule_df["DateStr"] = schedule_df["Date"].dt.strftime("%Y-%m-%d")
    schedule_df["Priority"] = schedule_df["task_type"].apply(
        lambda t: PRIORITY_ORDER.index(t) if t in PRIORITY_ORDER else len(PRIORITY_ORDER)
    )
    schedule_df["Task Type"] = schedule_df["task_type"].apply(
        lambda t: f"{TASK_TYPE_EMOJI.get(t, '')} {t}" if t else "❓"
    )
    schedule_df = schedule_df.sort_values(by=["Date", "Priority", "clean_task"])

    output_md = []

    if include_heading and heading_text:
        output_md.append(f"## {heading_text}\n")

    for date, group in schedule_df.groupby("Date"):
        day_str = date.strftime("%A, %Y-%m-%d")
        output_md.append(f"### 📅 {day_str}\n")

        header_cols = ["📜 Task", "🔖 Type"]
        if include_priority:
            header_cols.insert(1, "🔢 Priority")
        output_md.append("| " + " | ".join(header_cols) + " |")
        output_md.append("|" + " --- |" * len(header_cols))

        for _, row in group.iterrows():
            values = [row.get("clean_task", ""), row.get("Task Type", "")]
            if include_priority:
                values.insert(1, str(row.get("Priority", "")))
            output_md.append("| " + " | ".join(values) + " |")

        output_md.append("")

    return "\n".join(output_md)

def add_table_from_schedule_to_html(schedule_df: pd.DataFrame, section: str, include_priority: bool = True, include_heading: bool = False, heading_text: Optional[str] = None) -> str:
    if schedule_df.empty:
        return "<p><em>No schedule data available.</em></p>"

    schedule_df = schedule_df.copy()
    schedule_df["Date"] = pd.to_datetime(schedule_df["Date"], errors="coerce")
    schedule_df["Day"] = schedule_df["Date"].dt.strftime("%A")
    schedule_df["DateStr"] = schedule_df["Date"].dt.strftime("%Y-%m-%d")
    schedule_df["Priority"] = schedule_df["task_type"].apply(
        lambda t: PRIORITY_ORDER.index(t) if t in PRIORITY_ORDER else len(PRIORITY_ORDER)
    )
    schedule_df["Task Type"] = schedule_df["task_type"].apply(
        lambda t: f"{TASK_TYPE_EMOJI.get(t, '')} {t}" if t else "❓"
    )
    schedule_df = schedule_df.sort_values(by=["Date", "Priority", "clean_task"])

    html_output = []

    if include_heading and heading_text:
        html_output.append(f"<h2>{heading_text}</h2>")

    for date, group in schedule_df.groupby("Date"):
        day_str = date.strftime("%A, %Y-%m-%d")
        html_output.append(f"<h3>📅 {day_str}</h3>")
        html_output.append("<table border='1' cellspacing='0' cellpadding='4' style='border-collapse: collapse;'>")

        columns = ["📜 Task", "🔖 Type"]
        if include_priority:
            columns.insert(1, "🔢 Priority")
        html_output.append("<thead><tr>" + "".join(f"<th>{col}</th>" for col in columns) + "</tr></thead>")

        html_output.append("<tbody>")
        for _, row in group.iterrows():
            cells = [row.get("clean_task", ""), row.get("Task Type", "")]
            if include_priority:
                cells.insert(1, str(row.get("Priority", "")))
            html_output.append("<tr>" + "".join(f"<td>{str(cell)}</td>" for cell in cells) + "</tr>")
        html_output.append("</tbody></table><br>")

    return "\n".join(html_output)

def add_table_from_schedule_to_html(schedule_df: pd.DataFrame, section: str, include_priority: bool = True, include_heading: bool = False, heading_text: Optional[str] = None) -> str:
    if schedule_df.empty:
        return "<p><em>No schedule data available.</em></p>"

    schedule_df = schedule_df.copy()
    schedule_df["Date"] = pd.to_datetime(schedule_df["Date"], errors="coerce")
    schedule_df["Day"] = schedule_df["Date"].dt.strftime("%A")
    schedule_df["DateStr"] = schedule_df["Date"].dt.strftime("%Y-%m-%d")
    schedule_df["Priority"] = schedule_df["task_type"].apply(
        lambda t: PRIORITY_ORDER.index(t) if t in PRIORITY_ORDER else len(PRIORITY_ORDER)
    )
    schedule_df["Task Type"] = schedule_df["task_type"].apply(
        lambda t: f"{TASK_TYPE_EMOJI.get(t, '')} {t}" if t else "❓"
    )
    schedule_df = schedule_df.sort_values(by=["Date", "Priority", "clean_task"])

    html_output = []

    if include_heading and heading_text:
        html_output.append(f"<h2>{heading_text}</h2>")

    for date, group in schedule_df.groupby("Date"):
        day_str = date.strftime("%A, %Y-%m-%d")
        html_output.append(f"<h3>📅 {day_str}</h3>")
        html_output.append("<table border='1' cellspacing='0' cellpadding='4' style='border-collapse: collapse;'>")

        columns = ["📜 Task", "🔖 Type"]
        if include_priority:
            columns.insert(1, "🔢 Priority")
        html_output.append("<thead><tr>" + "".join(f"<th>{col}</th>" for col in columns) + "</tr></thead>")

        html_output.append("<tbody>")
        for _, row in group.iterrows():
            cells = [row.get("clean_task", ""), row.get("Task Type", "")]
            if include_priority:
                cells.insert(1, str(row.get("Priority", "")))
            html_output.append("<tr>" + "".join(f"<td>{str(cell)}</td>" for cell in cells) + "</tr>")
        html_output.append("</tbody></table><br>")

    return "\n".join(html_output)

def count_tokens(text: str, model: str = "gpt-3.5-turbo") -> int:
    try:
        enc = tiktoken.encoding_for_model(model)
    except KeyError:
        enc = tiktoken.get_encoding("cl100k_base")  # fallback
    return len(enc.encode(text))


def generate_llm_responses(blocks: List[str], api_key: str, model: str, debug: bool) -> List[str]:
    client = Mistral(api_key=api_key)
    markdown_output = []

    # ✅ Global Assertion: ensure all blocks are valid non-empty strings
    assert all(isinstance(b, str) and b.strip() for b in blocks), \
        "❌ One or more blocks are not valid non-empty strings."

    if debug:
        st.markdown("## 🧪 Block Sanity Check")
        st.markdown(f"- Total blocks: `{len(blocks)}`")
        for i, b in enumerate(blocks):
            token_count = count_tokens(b, model=model) if isinstance(b, str) else "N/A"
            st.text(f"Block {i+1}: type={type(b).__name__}, len={len(b.strip()) if isinstance(b, str) else 'N/A'}, tokens={token_count}")

    for i, block in enumerate(blocks):
        cleaned_block = block.strip()

        if debug:
            st.markdown(f"### 🧱 Prompt Block {i+1}")
            st.code(cleaned_block, language="markdown")
            st.text(f"🔢 Token count: {count_tokens(cleaned_block, model=model)}")
            st.markdown("### 📤 Block Sent to LLM")
            st.code(cleaned_block, language="markdown")

        try:
            st.info("⚙️ Calling LLM...")
            start_time = time.time()
            st.write("⏱️ LLM request started...")
        
            completion = client.chat.complete(
                model="mistral-small-latest",
                messages=[UserMessage(content=cleaned_block)],
                max_tokens=2048,
                temperature=0.5,
            )
            duration = time.time() - start_time
            st.write(f"✅ LLM response received in {duration:.2f} seconds")

            # ✅ Optional: debug raw object
            if debug:
                try:
                    st.markdown("### 🔎 Raw Completion Object (model_dump)")
                    st.json(completion.model_dump())
                except Exception:
                    try:
                        st.markdown("### 🔎 Raw Completion Object (dict fallback)")
                        st.json(completion.__dict__)
                    except Exception as e:
                        st.warning(f"⚠️ Could not log raw completion: {e}")
 
            response_text = completion.choices[0].message.content.strip()
            markdown_output.append(response_text)
        
        except Exception as e:
            error_msg = f"❌ LLM error on block {i+1}: {e}"
            st.error(error_msg)
            st.write (f"[ERROR] Block {i+1} failed: {e}")
            markdown_output.append(error_msg)

    return markdown_output


def generate_docx_from_prompt_blocks(
    blocks: list[str],
    section: str,
    doc_heading: Optional[str],
    schedule_sources: dict[str, str],
    api_key: Optional[str] = None,
    model: Optional[str] = None,
    use_llm: bool = False,
    debug: bool = False,
    include_priority: bool = True,
    insert_main_heading: bool = False,
    include_heading: bool = False,
) -> tuple[io.BytesIO, str, str]:

    debug_warnings = []

    if debug:
        st.markdown("### 🔧 Running generate_docx_from_prompt_blocks")
        st.write("📎 Section:", section)
        st.write("📊 Using LLM:", use_llm)

    # Step 0: Ensure <<INSERT_MAIL_TRASH_SCHEDULE_TABLE>> is present
    schedule_placeholder = "<<INSERT_MAIL_TRASH_SCHEDULE_TABLE>>"
    if not any(schedule_placeholder in block for block in blocks):
        blocks.append(f"### 📆 Combined Task Schedule\n{schedule_placeholder}")

    # Step 0.1: Add top-level heading if provided
    #if doc_heading:
     #   blocks.insert(0, f"# {doc_heading}")

    # Step 1: Generate markdown blocks (LLM or manual)
    markdown_blocks = []
    if use_llm:
        ##### debug
        if debug:
            st.markdown("### 🧪 Final Sanity Check Before LLM")
            for i, b in enumerate(blocks):
                st.markdown(f"#### 🧱 LLM Block {i+1}")
                st.markdown(f"**Type:** `{type(b).__name__}`")
                if not isinstance(b, str):
                    st.error(f"❌ Block {i+1} is NOT a string: {repr(b)}")
                    continue
                elif not b.strip():
                    st.warning(f"⚠️ Block {i+1} is blank or whitespace-only.")
                match = re.search(r"^#+\s+(.*)", b.strip(), re.MULTILINE)
                if match:
                    st.markdown(f"🏷️ Detected Title: `{match.group(1)}`")
                st.code(b, language="markdown")
                st.markdown(f"📝 Length: `{len(b.strip())}` characters")
                if len(b.strip()) < 50:
                    st.warning(f"⚠️ Block {i+1} is very short — may be a title-only block.")
        def is_placeholder_block(b: str) -> bool:
            return "<<INSERT_" in b and "SCHEDULE_TABLE>>" in b

        blocks_for_llm = [
            b for b in blocks if isinstance(b, str) and b.strip() and not is_placeholder_block(b)
        ]

        if debug:
            st.markdown("### 🔎 Prompt Blocks About to be Sent to LLM")
            for i, b in enumerate(blocks_for_llm):
                st.markdown(f"#### Block {i+1} (Type: `{type(b).__name__}`)")
                st.code(b, language="markdown")
            with open("debug_llm_blocks.json", "w") as f:
                json.dump(blocks_for_llm, f, indent=2)
        #### debug above
        try:
            markdown_blocks = generate_llm_responses(blocks_for_llm, api_key, model, debug)
         #### debug below  
            if debug:
                with open(f"debug_markdown_output_{section}.md", "w", encoding="utf-8") as f:
                    f.write("\n\n".join(markdown_blocks))
                st.success(f"📝 Debug markdown saved to `debug_markdown_output_{section}.md`")
        except Exception as e:
            st.error(f"❌ Error during LLM response generation: {e}")
            return "", f"LLM error: {e}"
    
    ### debug above
        else:
            markdown_blocks = [block for block in blocks if isinstance(block, str) and block.strip()]

    raw_markdown = "\n\n".join(markdown_blocks)
    
    # Step 2: Replace placeholders for markdown and HTML
    final_markdown = raw_markdown
    placeholder_to_df_map = {}
    for placeholder, df_key in schedule_sources.items():
        df = st.session_state.get(df_key)
        heading_text = SCHEDULE_HEADING_MAP.get(placeholder)
        if isinstance(df, pd.DataFrame) and not df.empty:
            markdown_table = add_table_from_schedule_to_markdown(
                df, section, include_priority=include_priority, include_heading=False, heading_text=heading_text
            )
            placeholder_to_df_map[placeholder] = markdown_table  # ✅ store resolved table
            final_markdown = final_markdown.replace(placeholder, markdown_table)
        else:
            warning = f"⚠️ No data found for `{df_key}`."
            placeholder_to_df_map[placeholder] = warning  # ✅ store fallback warning
            final_markdown = final_markdown.replace(placeholder, warning)
            if debug:
                st.markdown("### 🧪 Placeholder Table Replacements")
                for key, value in placeholder_to_df_map.items():
                    st.markdown(f"**{key}** →")
                    st.code(value, language="markdown")
                debug_warnings.append(f"Missing data for {df_key} (placeholder: {placeholder})")

    # Debug: render warnings
    if debug and debug_warnings:
        st.markdown("## ⚠️ Debug Warnings")
        for w in debug_warnings:
            st.markdown(f"- {w}")

    return final_markdown.strip(), debug_warnings

def preview_runbook_output(section: str, runbook_text: str, label: str = "📖 Preview Runbook"): ### depreciated ####
    """
    Debug/test-mode runbook preview with extra session context or metadata.

    Args:
        section (str): Used to determine which schedule data (if any) to insert.
        runbook_text (str): The raw LLM-generated markdown-style text.
        label (str): Button label to trigger the preview.
    """
    st.markdown(f"### {label} for `{section}`")

    if not runbook_text:
        st.warning("⚠️ No runbook content available to preview.")
        return

    # Optionally inject schedule table based on placeholder and section
    if "<<INSERT_SCHEDULE_TABLE>>" in runbook_text:
        df_key = "combined_home_schedule_df" if section == "home" else f"{section}_schedule_df"
        schedule_df = st.session_state.get(df_key)
        if isinstance(schedule_df, pd.DataFrame) and not schedule_df.empty:
            schedule_md = add_table_from_schedule_to_markdown(schedule_df)
            runbook_text = runbook_text.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)

    with st.expander("🧠 AI-Generated Runbook Preview", expanded=True):
        st.markdown(runbook_text, unsafe_allow_html=True)

def runbook_preview_dispatcher(
    section: str,
    runbook_text: str,
    mode: Literal["debug", "inline"] = "inline",
    show_schedule_snapshot: bool = False
):
    """
    Renders an AI-generated runbook preview using tabs to avoid nested expander errors.
    Includes optional schedule preview in a second tab if enabled and available.

    Args:
        section (str): The section the runbook is for.
        runbook_text (str): The raw markdown-style text from the LLM.
        mode (str): Either "debug" or "inline" to determine the label.
        show_schedule_snapshot (bool): Whether to include a tab with schedule preview.
    """
    if not runbook_text:
        st.warning("⚠️ No runbook content available to preview.")
        return

    # Lookup potential schedule
    df_key = "combined_home_schedule_df" if section == "home" else f"{section}_schedule_df"
    schedule_df = st.session_state.get(df_key)
    schedule_available = isinstance(schedule_df, pd.DataFrame) and not schedule_df.empty

    # Inject schedule if requested and placeholder is present
    if "<<INSERT_SCHEDULE_TABLE>>" in runbook_text and schedule_available:
        schedule_md = add_table_from_schedule_to_markdown(schedule_df)
        runbook_text = runbook_text.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)

    preview_label = "🧠 Runbook Preview (Debug)" if mode == "debug" else "📖 Runbook"
    show_tabs = show_schedule_snapshot and schedule_available

    # Always use tabs to avoid nested expander errors
    if show_tabs:
        tabs = st.tabs([preview_label, "📆 Schedule Snapshot"])
        with tabs[0]:
            st.markdown(runbook_text, unsafe_allow_html=True)
        with tabs[1]:
            st.markdown("### 📆 Scheduled Tasks Snapshot")
            st.dataframe(schedule_df)
    else:
        tabs = st.tabs([preview_label])
        with tabs[0]:
            st.markdown(runbook_text, unsafe_allow_html=True)

def render_runbook_preview_inline(
    section: str,
    runbook_text: str,
    schedule_df: Optional[pd.DataFrame],
    heading: str,
    timestamp: str
):
    """
    Renders a preview of the runbook and an optional schedule snapshot using tabs.
    """
    tabs = st.tabs(["📖 Runbook Text", "📊 Scheduled Tasks"])

    # 📝 Runbook Markdown View
    with tabs[0]:
        st.markdown(f"### {heading}")
        st.markdown(f"_Generated on {timestamp}_")

        if "<<INSERT_SCHEDULE_TABLE>>" in runbook_text and isinstance(schedule_df, pd.DataFrame):
            schedule_md = add_table_from_schedule_to_markdown(schedule_df)
            runbook_text = runbook_text.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)

        st.markdown(runbook_text)

        # 🧾 Optional raw preview for debugging
        with st.expander("🔍 View Raw Markdown"):
            st.code(runbook_text, language="markdown")

    # 📅 Schedule DataFrame Snapshot
    with tabs[1]:
        if schedule_df is not None and not schedule_df.empty:
            st.dataframe(schedule_df)
        else:
            st.info("ℹ️ No scheduled task snapshot available.")

def maybe_render_download(section: str) -> bool:
    """
    Render a minimal markdown-only runbook preview with a horizontal download + reset row.
    """
    text_key = f"{section}_runbook_text"
    preview_key = f"{section}_show_preview"

    runbook_text = st.session_state.get(text_key)

    if not runbook_text:
        return False

    if "<<" in runbook_text and ">>" in runbook_text:
        st.warning("⚠️ Some schedule placeholders may not have been replaced.")

    runbook_preview_dispatcher(
        section=section,
        runbook_text=runbook_text,
        mode="inline",
        show_schedule_snapshot=True
    )

    # 📥 Download + ♻️ Reset buttons in a row
    col1, col2 = st.columns([1, 1])

    with col1:
        st.download_button(
            label="📥 Download as Markdown (.md)",
            data=runbook_text,
            file_name=f"{section}_runbook.md",
            mime="text/markdown",
            key=f"Download_MD_{section}"
        )

    with col2:
        if st.button("♻️ Reset Cache", key=f"reset_{section}"):
            for k in [text_key, preview_key]:
                st.session_state.pop(k, None)
            st.success(f"🔄 Cleared runbook preview for `{section}`.")
            st.stop()

    return True


def maybe_generate_runbook(
    section: str,
    generator_fn: Callable[[], tuple[str, list[str]]],
    *,
    doc_heading: Optional[str] = None,
    button_label: str = "📥 Generate Runbook",
    filename: Optional[str] = None
):
    """
    Only generate runbook if requested by button and not already cached.

    Args:
        section (str): e.g., 'emergency_kit'
        generator_fn: returns (buffer, markdown_text, html_output)
        doc_heading (str): Optional heading for display.
        button_label (str): Button text.
        filename (str): Optional download filename.
    """
    generate_key = f"{section}_generate"
    ready_key = f"{section}_runbook_ready"
    text_key = f"{section}_runbook_text"
    debug = st.session_state.get("enable_debug_mode", False)

    st.subheader("🎉 Reward")
    if st.button("📤 Generate Runbook", key=generate_key):
        st.session_state[ready_key] = False
        st.rerun()

    if not st.session_state.get(ready_key):
        # Not ready yet, try generating
        markdown_text, debug_warnings = generator_fn()
        debug = st.session_state.get("enable_debug_mode", False)
        if debug and debug_warnings:
            st.markdown("### ⚠️ Missing Schedule Warnings")
            for msg in debug_warnings:
                st.markdown(f"- {msg}", unsafe_allow_html=True)
        if not isinstance(markdown_text, str):
            st.error("❌ markdown_text is not a string!")
        elif not markdown_text.strip():
            st.warning("⚠️ Runbook generation did not produce any content.")

        # Store in session
        st.session_state[text_key] = markdown_text
        st.session_state[ready_key] = True

        if st.session_state.get("enable_debug_mode"):
            st.success(f"✅ Generated runbook for `{section}`")
            st.write("📋 Final Markdown Output:", markdown_text[:500])  # Preview
            #st.write("📁 DOCX Buffer:", buffer)

    # Show download
    if st.session_state.get(ready_key):
        maybe_render_download(section=section)
        st.session_state.setdefault("level_progress", {})[section] = True
    else:
        st.info("ℹ️ Click the button above to generate your runbook.")

