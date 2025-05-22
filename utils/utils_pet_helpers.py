import streamlit as st
from mistralai import Mistral, UserMessage, SystemMessage
import csv
import io
from datetime import datetime, timedelta
from docx import Document
from collections import defaultdict
import json
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import pandas as pd
import re
from utils.data_access import (
    get_saved_pets_by_species,
)

def check_bingo(answers):
    for i in range(7):
        if all(answers[i][j] for j in range(7)) or all(answers[j][i] for j in range(7)):
            return True
    if all(answers[i][i] for i in range(7)) or all(answers[i][6 - i] for i in range(7)):
        return True
    return False

def flatten_answers_to_dict(questions, answers):
    return {
        questions[row * 7 + col]['label']: answers[row][col]
        for row in range(7) for col in range(7)
        if answers[row][col].strip()
    }

def get_pet_name(answers):
    try:
        return answers[0][0].strip() or "UnnamedPet"
    except:
        return "UnnamedPet"

def convert_to_csv(data_list):
    if not data_list:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=data_list[0].keys())
    writer.writeheader()
    writer.writerows(data_list)
    return output.getvalue()

def export_all_pets_to_docx(saved_pets, species, all_question_metadata):
    doc = Document()
    doc.add_heading(f"{species.capitalize()} Care Report", 0)
    for i, pet in enumerate(saved_pets, 1):
        name = pet.get("🐕 Pet Name", f"{species.capitalize()} #{i}")
        doc.add_heading(f"{i}. {name}", level=1)
        categories = {}
        for question, answer in pet.items():
            for q in all_question_metadata[species]:
                if q["label"] == question:
                    cat = q["category"]
                    categories.setdefault(cat, []).append((q["label"], answer))
                    break
        for cat, qas in categories.items():
            doc.add_heading(cat, level=2)
            for label, answer in qas:
                doc.add_paragraph(f"{label}: {answer}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def load_pet_for_edit(index, state_prefix, questions):
    selected = st.session_state[f"saved_{state_prefix}s"][index]
    new_answers = [['' for _ in range(7)] for _ in range(7)]
    for i in range(49):
        row, col = divmod(i, 7)
        label = questions[i]['label']
        new_answers[row][col] = selected.get(label, "")
    st.session_state[f"answers_{state_prefix}"] = new_answers
    st.session_state[f"editing_index_{state_prefix}"] = index
    st.rerun()

import re
import pandas as pd

def extract_pet_scheduled_tasks_grouped(questions, saved_pets, valid_dates):
    """
    Extracts pet care tasks based on valid_dates, labels, and user answers.
    Supports daily, weekly, every X weeks, monthly, and one-time tasks.
    Returns:
        - DataFrame with columns: Day, Pet, Task, Tag, Category
        - List of warning strings
    """
    schedule_rows = []
    added_tasks = set()
    warnings = []

    recurring_keywords = {
        "Feeding & Water": ["feed", "feeding schedule", "food", "treats", "water", "bowl"],
        "Grooming": ["brush", "bathing", "grooming", "nail", "teeth", "ear cleaning"],
        "Exercise / Walks": ["walk", "walking", "exercise"],
        "Medication": ["medication", "pill", "dose", "schedule"],
        "Behavior / Training": ["training", "commands", "behavior", "fear", "trigger"],
        "Play / Enrichment": ["play", "toys", "activities", "enrichment"],
        "Litter / Waste": ["litter", "waste", "box"],
        "Vet / Appointments": ["vet", "appointment", "check-up", "vaccination"],
        "Emergency / Contact": ["emergency", "contact", "phone", "email", "sitter", "walker"]
    }

    one_time_keywords = ["pill", "appointment", "check-up", "vaccination", "vet", "visit"]

    emoji_tags = {
        "[Daily]": "🔁 Daily",
        "[Weekly]": "📅 Weekly",
        "[One-Time]": "🗓️ One-Time",
        "[X-Weeks]": "↔️ Every X Weeks",
        "[Monthly]": "📆 Monthly"
    }

    weekday_to_int = {
        "Monday": 0, "Tuesday": 1, "Wednesday": 2, "Thursday": 3,
        "Friday": 4, "Saturday": 5, "Sunday": 6
    }

    def extract_week_interval(text):
        match = re.search(r"every (\d+) week", text.lower())
        if match:
            return int(match.group(1))
        if "biweekly" in text.lower():
            return 2
        return None

    def extract_weekday_mentions(text):
        return [day for day in weekday_to_int if day.lower() in text.lower()]

    def extract_dates(text):
        patterns = [
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}(?:,\s*\d{4})?",
            r"\b\d{1,2}/\d{1,2}(?:/\d{2,4})?"
        ]
        matches = []
        for pattern in patterns:
            matches += re.findall(pattern, text, flags=re.IGNORECASE)
        return matches

    def normalize_date(date_str):
        try:
            return pd.to_datetime(date_str, errors="coerce").date()
        except:
            return None

    def get_category(label):
        label_lower = label.lower()
        for category, keywords in recurring_keywords.items():
            if any(kw in label_lower for kw in keywords):
                return category
        return "Other"

    def get_pet_display_name(pet_dict):
        for k, v in pet_dict.items():
            if "name" in k.lower() and isinstance(v, str) and v.strip():
                return v.strip()
        return "Unnamed Pet"

    def should_exclude_label(label):
        exclusions = [
            "contact", "phone", "email", "sitter", "walker",
            "favorite", "toys", "play styles", "fear", "behavioral",
            "socialization", "appearance", "trainer", "photo", "description", "training goals",
            "training methods", "training challenges", "walking equipment", "walk behavior",
            "food brand/type your dog eats", "commands known"
        ]
        label_lower = label.lower()
        return any(kw in label_lower for kw in exclusions)

    def add_task(day, pet, task, tag, category):
        key = (day, pet, task)
        if key not in added_tasks:
            schedule_rows.append({
                "Day": day,
                "Pet": pet,
                "Task": task,
                "Tag": tag,
                "Category": category
            })
            added_tasks.add(key)

    for pet in saved_pets:
        pet_name = get_pet_display_name(pet)
        for q in questions:
            label = q["label"]
            answer = pet.get(label, "").strip()
            if not answer or should_exclude_label(label):
                continue

            task_text = f"{label} – {answer}"
            category = get_category(label)
            interval = extract_week_interval(answer)
            weekday_mentions = extract_weekday_mentions(answer)

            # Daily
            if (
                category != "Other"
                and not interval
                and not weekday_mentions
                and "monthly" not in answer.lower()
            ):
                for d in valid_dates:
                    add_task(d.strftime('%A'), pet_name, task_text, emoji_tags["[Daily]"], category)

            # Every X weeks
            if interval:
                if weekday_mentions:
                    for wd in weekday_mentions:
                        weekday_int = weekday_to_int[wd]
                        matching_dates = [d for d in valid_dates if d.weekday() == weekday_int]
                        if matching_dates:
                            base_date = matching_dates[0]
                            for d in matching_dates:
                                if (d - base_date).days % (interval * 7) == 0:
                                    add_task(d.strftime('%A'), pet_name, task_text, f"↔️ Every {interval} Weeks", category)
                else:
                    base_date = valid_dates[0]
                    for d in valid_dates:
                        if (d - base_date).days % (interval * 7) == 0:
                            add_task(d.strftime('%A'), pet_name, task_text, f"↔️ Every {interval} Weeks", category)
                    warnings.append(
                        f"⚠️ '{label}' mentions every {interval} weeks but no weekday. Defaulted to {base_date.strftime('%A')}."
                    )

            # Weekly
            if weekday_mentions and not interval:
                for wd in weekday_mentions:
                    weekday_idx = weekday_to_int[wd]
                    for d in valid_dates:
                        if d.weekday() == weekday_idx:
                            add_task(d.strftime('%A'), pet_name, task_text, emoji_tags["[Weekly]"], category)

            # Monthly
            if "monthly" in answer.lower():
                base_date = pd.to_datetime(valid_dates[0])
                current_date = base_date
                while current_date.date() <= valid_dates[-1]:
                    if current_date.date() in valid_dates:
                        add_task(current_date.strftime('%A'), pet_name, task_text, emoji_tags["[Monthly]"], category)
                    current_date += pd.DateOffset(months=1)
                warnings.append(
                    f"⚠️ '{label}' mentions monthly frequency with no specific weekday. Scheduled starting {base_date.strftime('%Y-%m-%d')}."
                )

            # One-time
            for ds in extract_dates(answer):
                parsed_date = normalize_date(ds)
                if parsed_date and parsed_date in valid_dates:
                    add_task(parsed_date.strftime('%A'), pet_name, task_text, emoji_tags["[One-Time]"], category)

    df = pd.DataFrame(schedule_rows)
    df = df.sort_values(by=["Pet", "Day", "Category", "Tag", "Task"]).reset_index(drop=True)
    return df, warnings

def get_pet_display_name(pet_dict):
    """Try to retrieve a pet's name from any key that looks like 'Name'."""
    for k in pet_dict:
        if "name" in k.lower():
            val = pet_dict[k].strip()
            if val:
                return val
    return "Unnamed Pet"
