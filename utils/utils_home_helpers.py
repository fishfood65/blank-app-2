import streamlit as st
import json
import csv
import io
from datetime import datetime, timedelta, date
from docx import Document
import pandas as pd
import plotly.express as px
from uuid import uuid4
import re
from collections import defaultdict
from docx.shared import Inches
from PIL import Image
import io
from .common_helpers import get_schedule_utils
from .input_tracker import capture_input, log_provider_result

def get_home_inputs():
    city = capture_input("City", st.text_input, "Home Basics")
    zip_code = capture_input("ZIP Code", st.text_input, "Home Basics")
    internet_provider = capture_input("Internet Provider", st.text_input, "Home Basics")
    st.session_state.city = city
    st.session_state.zip_code = zip_code
    st.session_state.internet_provider = internet_provider
    return city, zip_code, internet_provider


def get_corrected_providers(results):
    updated = {}

    label_to_key = {
    "Electricity": "electricity",
    "Natural Gas": "natural_gas",
    "Water": "water"
    }

    for label in ["Electricity", "Natural Gas", "Water"]:
        key = label_to_key[label]
        current_value = results.get(key, "") # Use get() to avoid KeyError

        correct_flag = st.checkbox(f"Correct {label} Provider", value=False)
        corrected = st.text_input(
            f"{label} Provider",
            value=current_value,
            disabled=not correct_flag
        )
        if correct_flag and corrected != current_value:
            log_provider_result(label, corrected)
            st.session_state[f"{key}_provider"] = corrected
        updated[key] = corrected if correct_flag else current_value
    return updated

def update_session_state_with_providers(updated):
    st.session_state["utility_providers"] = updated
    for key, value in updated.items():
        st.session_state[f"{key}_provider"] = value

def add_home_schedule_to_docx(doc, schedule_df):
    """
    Adds a grouped home schedule (Source → Date) to DOCX with embedded images inside task cells.

    Args:
        doc (Document): python-docx Document.
        schedule_df (pd.DataFrame): DataFrame with Task, Tag, Date, Source.
    """

    if schedule_df.empty:
        return

    # Set all categories to 'home'
    schedule_df["Category"] = "home"
    schedule_df["Date"] = pd.to_datetime(schedule_df["Date"], errors="coerce")
    schedule_df = schedule_df.sort_values(by=["Source", "Date", "Tag", "Task"])

    # Build image map from session
    image_map = {}
    if "trash_images" in st.session_state:
        for label, img_bytes in st.session_state["trash_images"].items():
            if img_bytes:
                keyword = label.replace(" Image", "").strip().lower()
                image_map[keyword] = img_bytes

    # Begin DOCX layout
    doc.add_page_break()
    doc.add_heading("📆 Home Maintenance Schedule", level=1)

    for source, source_group in schedule_df.groupby("Source"):
        doc.add_heading(f"🗂️ {source}", level=2)

        for date, date_group in source_group.groupby("Date"):
            day = date.strftime("%A")
            date_str = date.strftime("%Y-%m-%d")
            doc.add_heading(f"{day}, {date_str}", level=3)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Task"
            hdr_cells[1].text = "Tag"
            hdr_cells[2].text = "Category"

            for _, row in date_group.iterrows():
                task_text = str(row["Task"])
                tag = str(row["Tag"])
                category = str(row["Category"])
                task_lower = task_text.lower()

                cells = table.add_row().cells

                # Write text first
                paragraph = cells[0].paragraphs[0]
                run = paragraph.add_run(task_text)

                # Match image to task
                for keyword, image_bytes in image_map.items():
                    if keyword in task_lower:
                        try:
                            image_stream = io.BytesIO(image_bytes)
                            image = Image.open(image_stream)
                            image.thumbnail((500, 500))  # Resize if needed

                            resized_stream = io.BytesIO()
                            image.save(resized_stream, format="PNG")
                            resized_stream.seek(0)

                            paragraph.add_run().add_picture(resized_stream, width=Inches(2.5))
                        except Exception as e:
                            paragraph.add_run(f"\n⚠️ Failed to embed image: {e}")
                        break  # One image per task

                cells[1].text = tag
                cells[2].text = category

        doc.add_paragraph("")  # spacing

def generate_convenience_tasks(section_data: dict) -> list[dict]:
    """
    Generates task dictionaries from convenience_seeker section inputs.
    Each dict includes fields required for schedule generation.
    """
    tasks = []
    for service, answers in section_data.items():
        service_name = service.replace(" ", "_").lower()

        company = answers.get(f"{service} Company Name", "")
        phone = answers.get(f"{service} Company Phone Number", "")
        freq = answers.get(f"{service} Frequency", "")
        day = answers.get(f"{service} Day of the Week", "")
        access = answers.get(f"Access Method for {service}", "")
        postproc = answers.get(f"Post-{service} Procedures", "")
        verify = answers.get(f"{service} Crew Identity Verification", "")

        if freq and day:
            task = {
                "Task": f"{service} Service",
                "Category": "home",
                "Area": "Quality-Oriented Services",
                "Source": "convenience_seeker",
                "Tag": service_name,
                "Date": None,  # To be filled in later when expanded to actual calendar dates
                "Day": day,
                "Frequency": freq,
                "Details": {
                    "Company": company,
                    "Phone": phone,
                    "Access": access,
                    "PostProcedure": postproc,
                    "Verification": verify
                }
            }
            tasks.append(task)
    return tasks

