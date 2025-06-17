from io import BytesIO
import io
import re
from typing import Optional
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime
from collections import OrderedDict
from utils.data_helpers import get_answer, load_kit_items_config, safe_strip

KIT_CONFIG = load_kit_items_config()
UTILITY_KIT_RECS = KIT_CONFIG.get("per_utility", {})

def format_provider_markdown(providers: dict, version: str = "v1.0") -> str:
    """
    Returns markdown-formatted utility provider info using tables and icons.
    Adds support for utility-specific extra fields like 'Outage Support' for Internet.
    """
    icon_map = {
        "electricity": "⚡",
        "natural_gas": "🔥",
        "water": "💧",
        "internet": "🌐"
    }

    # 🔹 Date and version heading
    today = datetime.today().strftime("%B %d, %Y")
    heading = (
        "# ⚡🔥💧🌐 Utility Providers Emergency Guide\n\n"
        f"📅 **Generated:** {today} | **Version:** {version}\n\n"
        "This guide includes a summary and detailed emergency info for each utility.\n"
    )

    # 🔹 Summary Table
    summary_lines = [
        "### 🗂️ Utility Provider Summary", "",
        "| Utility | Provider | Phone | Website |",
        "|---------|----------|-------|---------|"
    ]
    for key, info in providers.items():
        icon = icon_map.get(key, "📦")
        name = info.get("name", "⚠️ Not Available").strip()
        phone = info.get("contact_phone", "⚠️ Not Available").strip()
        website = info.get("contact_website", "").strip()
        website_display = f"[{website}]({website})" if website.startswith("http") else website or "⚠️ Not Available"

        summary_lines.append(
            f"| {icon} {key.replace('_', ' ').title()} | {name} | {phone} | {website_display} |"
        )

    # 🔹 Detailed Sections
    output = []

    for key, info in providers.items():
        name = info.get("name", "").strip()
        if not name or name.lower() in ["⚠️ not available", "n/a", "not found"]:
            continue

        icon = icon_map.get(key, "📦")
        label = key.replace("_", " ").title()
        section_title = f"### {icon} {label} – {name}"

        table_lines = ["| Field | Value |", "|-------|--------|"]

        def row(label, field, is_link=False, multiline=False): # Removed is_email=False
            value = info.get(field, "").strip()
            if not value or value.lower() in ["⚠️ not available", "n/a", "not found"]:
                return None
            if is_link:
                return f"| **{label}** | [{value}]({value}) |"
            #elif is_email and "@" in value:
            #   return f"| **{label}** | [{value}](mailto:{value}) |"
            elif multiline:
                value_md = value.replace("\n", "<br>")
                return f"| **{label}** | {value_md} |"
            else:
                return f"| **{label}** | {value} |"

        # Standard fields
        standard_fields = [
            ("Description", "description", False, False),
            ("Phone", "contact_phone", False, False),
            ("Website", "contact_website", True, False),
            #("Email", "contact_email", False, True),
            ("Address", "contact_address", False, False),
        ]

        for label_, field_, is_link_, multiline_ in standard_fields:
            line = row(label_, field_, is_link_, multiline_)
            if line:
                table_lines.append(line)

        section_block = section_title + "\n" + "\n".join(table_lines)

        # 🚨 Emergency Steps
        emergency = info.get("emergency_steps", "").strip()
        if emergency and emergency.lower() != "⚠️ not available":
            formatted = "\n".join([
                line.strip() if line.strip().startswith(("-", "•", "*")) else f"- {line.strip()}"
                for line in emergency.splitlines() if line.strip()
            ])
            section_block += "\n\n### 🚨 Emergency Instructions\n" + formatted

        # 🧾 Non-Emergency Tips
        tips = info.get("non_emergency_tips", "").strip()
        if tips and tips.lower() != "⚠️ not available":
            formatted_tips = "\n".join([
                line.strip() if line.strip().startswith(("-", "•", "*")) else f"- {line.strip()}"
                for line in tips.splitlines() if line.strip()
            ])
            section_block += "\n\n### 🧾 Non-Emergency Tips\n" + formatted_tips

        output.append(section_block)

    return "\n".join([heading] + summary_lines + ["\n---\n"] + output)

def render_recommended_items_for(utility: str) -> str:
    items = UTILITY_KIT_RECS.get(utility, [])
    if not items:
        return "_(none specified)_"
    return "\n".join(f"- {item}" for item in items)

from datetime import datetime

def generate_emergency_utilities_kit_markdown() -> str:
    from datetime import datetime

    config = load_kit_items_config()
    KIT_ITEMS = config.get("recommended_items", [])
    PER_UTILITY = config.get("per_utility", {})

    def render_recommended(utility_key):
        items = PER_UTILITY.get(utility_key, [])
        if not items:
            return "_(none specified)_"
        return "\n".join(f"- {item}" for item in items)

    # Pull values from session_state
    city = safe_strip(get_answer("City", section="utilities"))
    zip_code = safe_strip(get_answer("ZIP Code", section="utilities"))
    providers = st.session_state.get("corrected_utility_providers", {})

    electricity = safe_strip(providers.get("electricity", {}).get("name", "⚠️ Not provided"))
    gas = safe_strip(providers.get("natural_gas", {}).get("name", "⚠️ Not provided"))
    water = safe_strip(providers.get("water", {}).get("name", "⚠️ Not provided"))
    internet = safe_strip(providers.get("internet", {}).get("name", "⚠️ Not provided"))

    selected = st.session_state.get("homeowner_kit_stock", [])
    missing = st.session_state.get("not_selected_items", [])
    additional = st.session_state.get("additional_kit_items", "")
    matched = st.session_state.get("matched_additional_items", [])
    unmatched = st.session_state.get("unmatched_additional_items", [])

    # Emergency Kit section
    parts = ["# 🧰 Emergency Kit Summary"]

    if selected:
        parts.append("## ✅ Kit Items You Have:")
        parts.append("\n".join(f"- {item}" for item in selected))
    if missing:
        parts.append("## ⚠️ Missing Recommended Items:")
        parts.append("\n".join(f"- {item}" for item in missing))
    if matched:
        parts.append("## 🔁 Additional Items Matched to Kit List:")
        parts.append("\n".join(f"- {item}" for item in matched))
    if unmatched:
        parts.append("## ❓ Custom Items (Unmatched):")
        parts.append("\n".join(f"- {item}" for item in unmatched))
    if additional and not (matched or unmatched):
        parts.append("## ➕ Additional User-Entered Items:")
        parts.append(additional)

    # Utilities overview section
    parts.append("\n# 🏡 Emergency Utilities Overview")
    today = datetime.today().strftime("%B %d, %Y")
    parts.append(f"📅 **Generated:** {today} | **Location:** {city}, ZIP {zip_code}\n")
    parts.append("This guide includes a summary and detailed emergency info for each utility.\n")

    # Summary Table
    parts.append("### 🗂️ Utility Provider Summary\n")
    summary_lines = [
        "| Utility | Provider | Phone | Website |",
        "|---------|----------|-------|---------|"
    ]
    icon_map = {
        "electricity": "⚡",
        "natural_gas": "🔥",
        "water": "💧",
        "internet": "🌐"
    }

    for key, info in providers.items():
        icon = icon_map.get(key, "📦")
        name = info.get("name", "⚠️ Not Available").strip()
        phone = info.get("contact_phone", "⚠️ Not Available").strip()
        website = info.get("contact_website", "").strip()
        website_display = f"[{website}]({website})" if website.startswith("http") else website or "⚠️ Not Available"
        summary_lines.append(f"| {icon} {key.replace('_', ' ').title()} | {name} | {phone} | {website_display} |")

    parts.append("\n".join(summary_lines))
    parts.append("---")  # section break

    # Detailed sections
    def add_utility_block(label, key, provider_key):
        provider = providers.get(provider_key, {})
        name = provider.get("name", "").strip()
        if not name or "⚠️" in name:
            return

        parts.append(f"### {label} – {name}")

        table_lines = ["| Field | Value |", "|-------|--------|"]
        def row(label, field, is_link=False, multiline=False):
            value = provider.get(field, "").strip()
            if not value or value.lower() in ["⚠️ not available", "n/a", "not found"]:
                return None
            if is_link:
                return f"| **{label}** | [{value}]({value}) |"
            elif multiline:
                return f"| **{label}** | {value.replace(chr(10), '<br>')} |"
            else:
                return f"| **{label}** | {value} |"

        for label_, field_, is_link_ in [
            ("Description", "description", False),
            ("Phone", "contact_phone", False),
            ("Website", "contact_website", True),
            ("Address", "contact_address", False)
        ]:
            line = row(label_, field_, is_link_)
            if line:
                table_lines.append(line)

        parts.append("\n".join(table_lines))

        emergency = provider.get("emergency_steps", "").strip()
        if emergency and emergency.lower() != "⚠️ not available":
            # 🧹 Strip out any "Non-Emergency Tips" leakage from emergency_steps
            emergency = re.split(r"(?i)non[- ]?emergency tips[:\-]?", emergency)[0].strip()
            formatted = "\n".join([
                line.strip() if line.strip().startswith(("-", "•", "*")) else f"- {line.strip()}"
                for line in emergency.splitlines() if line.strip()
            ])
            parts.append("\n#### 🚨 Emergency Instructions\n" + formatted)

        tips = provider.get("non_emergency_tips", "").strip()
        if tips and tips.lower() != "⚠️ not available":
            formatted = "\n".join([
                line.strip() if line.strip().startswith(("-", "•", "*")) else f"- {line.strip()}"
                for line in tips.splitlines() if line.strip()
            ])
            parts.append("\n#### 🧾 Non-Emergency Tips\n" + formatted)

        parts.append("#### 🧰 Recommended Kit Items:")
        items = PER_UTILITY.get(provider_key, [])
        if items:
            parts.append("\n".join(f"- {item}" for item in items))
        else:
            parts.append("_(none specified)_")

    add_utility_block("⚡ Electricity", "electricity", "electricity")
    add_utility_block("🔥 Natural Gas", "natural_gas", "natural_gas")
    add_utility_block("💧 Water", "water", "water")
    add_utility_block("🌐 Internet", "internet", "internet")

    return "\n\n".join(parts + [""]).strip()

def add_hyperlink(paragraph, url, text=None):
    """
    Add a hyperlink to a paragraph, styled as blue and underlined.
    """
    if text is None:
        text = url

    # Create relationship ID
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create hyperlink tag
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create run element
    new_run = OxmlElement("w:r")

    # Create run properties with blue color and underline
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # Blue

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")  # Underline

    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)

    # Add the text
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return paragraph

def add_provider_summary_table(doc: Document, providers: dict):
    """
    Adds a summary table to the given Word doc listing all utility providers
    with utility type, name, phone number, and website (as hyperlink).
    """
    icon_map = {
        "electricity": "⚡",
        "natural_gas": "🔥",
        "water": "💧",
        "internet": "🌐"
    }

    # Add a heading for the summary
    doc.add_heading("🗂️ Utility Provider Summary", level=2)
    doc.add_paragraph("")  # spacer

    # Create a table with headers
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = True
    table.allow_autofit = True

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Utility"
    hdr_cells[1].text = "Provider Name"
    hdr_cells[2].text = "Phone"
    hdr_cells[3].text = "Website"

    # Add a row per provider
    for utility, info in providers.items():
        name = info.get("name", "").strip()
        phone = info.get("contact_phone", "").strip()
        website = info.get("contact_website", "").strip()

        if not name or name.lower() == "⚠️ not available":
            continue

        icon = icon_map.get(utility, "🔌")
        row_cells = table.add_row().cells
        row_cells[0].text = f"{icon} {utility.replace('_', ' ').title()}"
        row_cells[1].text = name
        row_cells[2].text = phone

        if website.startswith("http"):
            # Add hyperlink inside cell
            p = row_cells[3].paragraphs[0]
            add_hyperlink(p, website, text=website)
        else:
            row_cells[3].text = website


# Utility function to add a single provider section to an existing doc
def add_provider_section_to_docx(doc: Document, providers: dict):
    """
    Appends a styled section for each utility provider to the given Document object.
    Skips fields that are missing or empty.
    """
    icons = {
        "electricity": "⚡",
        "natural_gas": "🔥",
        "water": "💧",
        "internet": "🌐"
    }

    # Fields to render as a bullet list *below* the normal block
    multiline_fields = OrderedDict([
        ("emergency_steps", "🚨 Emergency Instructions"),
        ("non_emergency_tips", "🧾 Non-Emergency Tips")
    ])

    for utility, info in providers.items():
        name = info.get("name", "").strip()
        if not name or name.lower() == "⚠️ not available":
            continue

        icon = icons.get(utility, "🔌")
        doc.add_heading(f"{icon} {utility.replace('_', ' ').title()} – {name}", level=2)

        def add_field(label, key, bold=True):
            value = info.get(key, "").strip()
            if not value:
                return

            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            if bold:
                run.bold = True

            # Add hyperlink if Website or Email
            if key == "contact_website" and value.startswith("http"):
                add_hyperlink(para, value)
            elif key == "contact_email" and "@" in value:
                add_hyperlink(para, f"mailto:{value}", text=value)
            else:
                para.add_run(value)

            para.paragraph_format.space_after = Pt(6)

        add_field("Description", "description")
        add_field("Phone", "contact_phone")
       #add_field("Email", "contact_email")
        add_field("Website", "contact_website")
        add_field("Address", "contact_address")

        # ✅ Multiline fields rendered as bulleted lists
        for field_key, heading in multiline_fields.items():
            value = info.get(field_key, "").strip()
            if value and value.lower() not in ["⚠️ not available", "n/a"]:
                lines = [line.strip() for line in value.splitlines() if line.strip()]
                if lines:
                    doc.add_paragraph("")  # spacer
                    doc.add_heading(heading, level=3)
                    for line in lines:
                        # Strip leading symbols/numbers
                        clean_line = re.sub(r"^[-•*]|\d+[.)]\s*", "", line).strip()
                        if clean_line:
                            doc.add_paragraph(clean_line, style="List Bullet")

# Export wrapper
def export_provider_docx(providers: dict, output_path: str = None, version: str = "v1.0"):
    """
    Creates a full Document, adds all utility providers, and saves to disk.
    """
    doc = Document()
    doc.add_heading("🏡 Emergency Utilities Guide", level=1)

    # Date and version stamp
    today = datetime.today().strftime("%B %d, %Y")
    doc.add_paragraph(f"📅 Generated: {today} | Version: {version}")
    doc.add_paragraph("The guide contains contact and emergency information for each utility provider.")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_provider_summary_table(doc, providers)
    add_provider_section_to_docx(doc, providers)

    if output_path:
        doc.save(output_path)
        return None
    else:
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

def generate_emergency_utilities_kit_docx() -> BytesIO:
    config = load_kit_items_config()
    PER_UTILITY = config.get("per_utility", {})

    doc = Document()

    # --- Formatting helpers ---
    def add_heading(text, level=1):
        doc.add_heading(text, level=level)

    def add_paragraph(text, bold=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        if bold:
            run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    def add_bullet_list(items):
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    # --- Load session values ---
    providers = st.session_state.get("corrected_utility_providers", {})
    city = safe_strip(get_answer("City", section="utilities"))
    zip_code = safe_strip(get_answer("ZIP Code", section="utilities"))
    selected = st.session_state.get("homeowner_kit_stock", [])
    missing = st.session_state.get("not_selected_items", [])
    matched = st.session_state.get("matched_additional_items", [])
    unmatched = st.session_state.get("unmatched_additional_items", [])
    additional = st.session_state.get("additional_kit_items", "")

    # --- Emergency Kit Section ---
    add_heading("🧰 Emergency Kit Summary", level=1)

    if selected:
        add_heading("✅ Kit Items You Have:", level=2)
        add_bullet_list(selected)
    if missing:
        add_heading("⚠️ Missing Recommended Items:", level=2)
        add_bullet_list(missing)
    if matched:
        add_heading("🔁 Additional Items Matched to Kit List:", level=2)
        add_bullet_list(matched)
    if unmatched:
        add_heading("❓ Custom Items (Unmatched):", level=2)
        add_bullet_list(unmatched)
    if additional and not matched and not unmatched:
        add_heading("➕ Additional User-Entered Items:", level=2)
        add_paragraph(additional)

    # --- Emergency Utilities Overview ---
    add_heading("🏡 Emergency Utilities Overview", level=1)
    add_paragraph(f"📍 Location: {city}, ZIP {zip_code}")

    emoji_map = {
        "electricity": "⚡",
        "natural_gas": "🔥",
        "water": "💧",
        "internet": "🌐"
    }

    for utility_key in ["electricity", "natural_gas", "water", "internet"]:
        provider = providers.get(utility_key, {})
        name = provider.get("name", "").strip()
        if not name or name.lower() in ["⚠️ not available", "n/a", "not found"]:
            continue

        icon = emoji_map.get(utility_key, "🔌")
        add_heading(f"{icon} {utility_key.replace('_', ' ').title()} – {name}", level=2)

        add_paragraph("Recommended Kit Items:", bold=True)
        kit_items = PER_UTILITY.get(utility_key, [])
        if kit_items:
            add_bullet_list(kit_items)
        else:
            add_paragraph("_(none specified)_")

    # --- Provider Summary + Details ---
    if providers:
        add_heading("📡 Detailed Provider Contact Info", level=1)
        add_provider_summary_table(doc, providers)
        add_provider_section_to_docx(doc, providers)
    else:
        add_paragraph("⚠️ No confirmed provider data found.")

    # --- Final Output ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render_runbook_section_output(
    markdown_str: str,
    docx_bytes_io: Optional[BytesIO] = None,
    title: str = "Runbook Section",
    filename_prefix: str = "section",
    expand_preview: bool = False,
):
    """
    General-purpose viewer and exporter for runbook sections.
    """
    section_slug = filename_prefix.lower().replace(" ", "_")

    with st.expander(f"🧾 {title} Preview", expanded=expand_preview):
        if markdown_str:
            st.markdown(markdown_str, unsafe_allow_html=True)
        else:
            st.info(f"No markdown content available for {title}.")

    with st.expander(f"📁 Download {title}"):
        if docx_bytes_io:
            st.download_button(
                label=f"📄 Download {title} as DOCX",
                data=docx_bytes_io,
                file_name=f"{section_slug}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif markdown_str:
            st.download_button(
                label=f"⬇️ Download {title} as Markdown",
                data=markdown_str,
                file_name=f"{section_slug}.md",
                mime="text/markdown"
            )
        else:
            st.info(f"⚠️ No downloadable content for {title}.")
