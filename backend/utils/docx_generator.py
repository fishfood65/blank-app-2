from docx import Document
from io import BytesIO

def generate_docx(section: str, inputs: dict) -> BytesIO:
    doc = Document()
    doc.add_heading(f"{section.title()} Runbook", 0)

    for key, val in inputs.items():
        doc.add_heading(key, level=2)
        doc.add_paragraph(val or "⚠️ Not Provided")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
