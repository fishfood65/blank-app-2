from config.sections import (SECTION_METADATA, LLM_SECTIONS)
from utils.prompt_block_utils import generate_all_prompt_blocks
import streamlit as st
import re
from mistralai import Mistral, UserMessage, SystemMessage
from dotenv import load_dotenv
import os
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.text.run import Run
import re
import time
from PIL import Image
import io
import uuid
import json
from docx.shared import Inches, Pt
from typing import List, Tuple, Optional, Union
from utils.common_helpers import (
    extract_all_trash_tasks_grouped, 
    extract_grouped_mail_task, 
    generate_flat_home_schedule_markdown,
    switch_section,
    get_schedule_placeholder_mapping,
    debug_saved_schedule_dfs
)
from prompts.templates import(
    utility_provider_lookup_prompt
)
from utils.data_helpers import (
    capture_input, 
    flatten_answers_to_dict, 
    get_answer, 
    extract_and_log_providers, 
    log_provider_result, 
    preview_input_data, 
    check_missing_utility_inputs, 
    export_input_data_as_csv, 
    render_lock_toggle,
    daterange,
    get_filtered_dates,
    select_runbook_date_range,
    register_task_input,
    extract_providers_from_text,
)
from utils.runbook_generator_helpers import (
    maybe_render_download,
    maybe_generate_runbook,
    generate_docx_from_prompt_blocks,
    add_table_from_schedule,
    add_table_from_schedule_to_markdown
)
from utils.task_schedule_utils_updated import (
    extract_and_schedule_all_tasks,
    get_schedule_utils,
    generate_flat_home_schedule_markdown,
    save_task_schedules_by_type,
    extract_unscheduled_tasks_from_inputs_with_category,
    load_label_map,
    normalize_label
)
from utils.preview_helpers import (
    display_enriched_task_preview, 
    edit_button_redirect,
    get_active_section_label
)

def preview_runbook_output(runbook_text: str, label: str = "📖 Preview Runbook"):
    """
    Shows an expandable markdown preview of the runbook text when a button is clicked.

    Args:
        runbook_text (str): The raw LLM-generated markdown-style text.
        label (str): Button label to trigger the preview.
    """
    # use the following code with home_app_05_23_modified.py
    #if not runbook_text:
    #    st.warning("⚠️ No runbook content available to preview.")
    #    return

    #if st.button(label):
    #    with st.expander("🧠 AI-Generated Runbook Preview", expanded=True):
    #        st.markdown(runbook_text)
    if not runbook_text:
        st.warning("⚠️ No runbook content available to preview.")
        return

    if st.button(label):
        # Optionally add schedule preview
        schedule_df = st.session_state.get("combined_home_schedule_df")
        if isinstance(schedule_df, pd.DataFrame) and not schedule_df.empty:
            schedule_md = add_table_from_schedule_to_markdown(schedule_df)
            runbook_text = runbook_text.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)

        with st.expander("🧠 AI-Generated Runbook Preview", expanded=True):
            st.markdown(runbook_text, unsafe_allow_html=True)

def render_prompt_preview(missing: list, section: str = "home"):
    confirmed = st.session_state.get(f"{section}_user_confirmation", False)

    with st.expander("🧠 AI Prompt Preview (Optional)", expanded=True):
        if missing:
            st.warning(f"⚠️ Cannot generate prompt. Missing: {', '.join(missing)}")
            return

        if not confirmed:
            st.info("☕️ Please check the box to confirm AI prompt generation.")
            return

        prompt = st.session_state.get("generated_prompt", "")
        prompt_blocks = st.session_state.get("prompt_blocks", [])
        schedule_md = st.session_state.get("home_schedule_markdown", "_No schedule available._")

        if not prompt:
            st.warning("⚠️ Prompt not generated yet.")
            return

        # Build combined preview by inserting schedule into the final prompt block
        if prompt_blocks:
            full_preview = "\n\n".join(prompt_blocks)
            full_preview = full_preview.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)
        else:
            full_preview = prompt.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)

        # Display the formatted full prompt preview
        st.markdown(full_preview)
        if section == "mail_trash_handling":
            st.markdown("### 📋 Schedule Preview")
            st.markdown(st.session_state.get("home_schedule_markdown", "_No schedule available._"))
        st.success("✅ Prompt ready! This is what will be sent to the LLM.")

def generate_docx_from_split_prompts( ### depreciate 
    prompts: List[str],
    api_key: str,
    *,
    section_titles: Optional[List[str]] = None,
    model: str = "mistral-small-latest",
    doc_heading: str = "Runbook",
    temperature: float = 0.5,
    max_tokens: int = 2048,
    debug: bool = False
) -> Tuple[io.BytesIO, str]:

    if not prompts or not isinstance(prompts, list):
        raise ValueError("🚫 prompts must be a non-empty list of strings.")

    prompts = [p.strip() for p in prompts if p.strip()]
    if not prompts:
        raise ValueError("🚫 All prompts were empty after stripping.")

    combined_output = []

    for i, prompt in enumerate(prompts):
        try:
            if debug:
                st.markdown(f"### 🧾 Prompt Block {i+1}")
                st.code(prompt, language="markdown")

            client = Mistral(api_key=api_key)
            completion = client.chat.complete(
                model=model,
                messages=[SystemMessage(content=prompt)],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            response_text = completion.choices[0].message.content

            if debug:
                st.markdown(f"### 🧾 Raw LLM Response [Block {i + 1}]")
                st.code(response_text, language="markdown")

            title = section_titles[i].strip() if section_titles and i < len(section_titles) else ""
            section_label = f"### {title}" if title else ""
            output_block = f"{section_label}\n\n{response_text.strip()}" if section_label else response_text.strip()
            combined_output.append(output_block)

        except Exception as e:
            st.error(f"❌ Error processing prompt {i + 1}: {e}")
            continue

    full_text = "\n\n".join(combined_output)
    if debug:
        st.markdown("### 🔍 Full Text Passed to DOCX Builder")
        st.code(full_text, language="markdown")

    doc = Document()
    doc.add_heading(doc_heading, 0)
    lines = full_text.splitlines()

    if debug:
        st.markdown("### 🧾 Lines Sent to DOCX")
        st.code("\n".join(lines), language="markdown")

    schedule_sources = get_schedule_placeholder_mapping()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Table placeholders
        if line in schedule_sources:
            df_key = schedule_sources[line]
            schedule_df = st.session_state.get(df_key)

            doc.add_paragraph("")  # spacing
            if isinstance(schedule_df, pd.DataFrame) and not schedule_df.empty:
                if debug:
                    doc.add_paragraph(f"✅ [DEBUG] Inserted table for {line}")
                add_table_from_schedule(doc, schedule_df)
            else:
                doc.add_paragraph("_No schedule available._")
                if debug:
                    st.warning(f"⚠️ No data found for placeholder: {line}")
            doc.add_paragraph("")
            continue

        # Markdown-to-Word formatting
        if line.startswith("##### "):
            doc.add_heading(line[6:].strip(), level=4)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            para = doc.add_paragraph()
            cursor = 0
            for match in re.finditer(r"(\*\*.*?\*\*)", line):
                start, end = match.span()
                if start > cursor:
                    para.add_run(line[cursor:start])
                para.add_run(match.group(1)[2:-2]).bold = True
                cursor = end
            if cursor < len(line):
                para.add_run(line[cursor:])
            para.style.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, full_text

def mail_trash_runbook_prompt(debug_key="trash_info_debug") -> list:
    """
    Returns a list of smaller prompt blocks instead of one large prompt string.
    Suitable for passing into multi-step LLM document generation.
    """
    input_data = st.session_state.get("input_data", {})
    merged_entries = input_data.get("mail_trash_handling", [])

    if merged_entries:
        mail_entries = [e for e in merged_entries if "mail" in str(e.get("section", "")).lower()]
        trash_entries = [e for e in merged_entries if "trash" in str(e.get("section", "")).lower()]
    else:
        mail_entries = input_data.get("mail") or input_data.get("Mail & Packages", [])
        trash_entries = input_data.get("Trash Handling", []) or input_data.get("trash_handling", [])

    mail_info = {entry["question"]: entry["answer"] for entry in mail_entries}
    trash_info = {entry["question"]: entry["answer"] for entry in trash_entries}

    def safe_line(label, value):
        if value and str(value).strip().lower() != "no":
            return f"- **{label}**: {value}"
        return None

    def safe_yes_no(label, flag, detail_label, detail_value):
        if flag:
            return f"- **{label}**: Yes\n  - **{detail_label}**: {detail_value or 'N/A'}"
        return ""

    mail_block = "\n".join(filter(None, [
        safe_line("Mailbox Location", mail_info.get("\ud83d\udccd Mailbox Location")),
        safe_line("Mailbox Key Info", mail_info.get("\ud83d\udd11 Mailbox Key (Optional)")),
        safe_line("Pick-Up Schedule", mail_info.get("\ud83d\udcc6 Mail Pick-Up Schedule")),
        safe_line("Mail Sorting Instructions", mail_info.get("\ud83d\udce5 What to Do with the Mail")),
        safe_line("Delivery Packages", mail_info.get("\ud83d\udce6 Packages")),
    ]))

    indoor_block = "\n".join(filter(None, [
        safe_line("Kitchen Trash", trash_info.get("Kitchen Trash Bin Location, Emptying Schedule and Replacement Trash Bags")),
        safe_line("Bathroom Trash", trash_info.get("Bathroom Trash Bin Emptying Schedule and Replacement Trash Bags")),
        safe_line("Other Rooms Trash", trash_info.get("Other Room Trash Bin Emptying Schedule and Replacement Trash Bags")),
    ]))

    outdoor_lines = [
        safe_line("Please take the bins", trash_info.get("Instructions for Placing and Returning Outdoor Bins")),
        safe_line("Bins Description", trash_info.get("What the Outdoor Trash Bins Look Like")),
        safe_line("Location", trash_info.get("Specific Location or Instructions for Outdoor Bins")),
    ]

    outdoor_image_tags = []
    if "trash_images" in st.session_state:
        for label in ["Outdoor Bin Image", "Recycling Bin Image"]:
            if st.session_state.trash_images.get(label):
                filename = st.session_state.trash_images[label]
                outdoor_image_tags.append(f'<img src="{filename}" alt="{label}" width="300"/>')

    outdoor_block = "\n".join(filter(None, outdoor_lines + outdoor_image_tags))

    collection_block = "\n".join(filter(None, [
        safe_line("Garbage Pickup", f"{trash_info.get('Garbage Pickup Day', '')}, {trash_info.get('Garbage Pickup Time', '')}".strip(", ")),
        safe_line("Recycling Pickup", f"{trash_info.get('Recycling Pickup Day', '')}, {trash_info.get('Recycling Pickup Time', '')}".strip(", ")),
    ]))

    composting_used = trash_info.get("Compost Used", "").strip().lower() == "yes"
    composting_block = safe_yes_no("Composting Used", composting_used, "Compost Instructions", trash_info.get("Compost Instructions"))

    common_disposal_used = trash_info.get("Common Disposal Used", "").strip().lower() == "yes"
    common_disposal_block = safe_yes_no("Common Disposal Area Used", common_disposal_used, "Instructions", trash_info.get("Common Disposal Area Instructions"))

    wm_block = "\n".join(filter(None, [
        safe_line("Company Name", trash_info.get("Waste Management Company Name")),
        safe_line("Phone", trash_info.get("Contact Phone Number")),
        safe_line("When to Contact", trash_info.get("When to Contact")),
    ]))

    schedule_md = st.session_state.get("home_schedule_markdown", "_Schedule missing._")

    # Final prompt chunks
    return [
        """
You are an expert assistant generating a Mail Run Book. Compose a clear and concise guide for house sitters.

## 📬 Mail Handling Instructions

{mail_block}
""".strip().format(mail_block=mail_block),
        """
You are an expert assistant describing indoor trash handling instructions.

### Indoor Trash

{indoor_block}
""".strip().format(indoor_block=indoor_block),
        """
You are an expert assistant describing outdoor trash and bin logistics.

### Outdoor Bins

{outdoor_block}
""".strip().format(outdoor_block=outdoor_block),
        """
### Collection Schedule

{collection_block}
""".strip().format(collection_block=collection_block),
        """
### Composting

{composting_block}
""".strip().format(composting_block=composting_block),
        """
### Common Disposal Area

{common_disposal_block}
""".strip().format(common_disposal_block=common_disposal_block),
        """
### Waste Management Contact

{wm_block}
""".strip().format(wm_block=wm_block),
        """
## 📆 Mail & Trash Pickup Schedule

{schedule_md}
""".strip().format(schedule_md=schedule_md),
    ]

def emergency_kit_utilities_runbook_prompt():
    """
    Generate a markdown-formatted emergency runbook prompt using user answers and utility data.
    """
    city = get_answer("City", "Home Basics") or ""
    zip_code = get_answer("ZIP Code", "Home Basics") or ""
    internet_provider_name = get_answer("Internet Provider", "Home Basics") or ""
    emergency_kit_status = get_answer("Do you have an Emergency Kit?", "Emergency Kit") or "No"
    emergency_kit_location = get_answer("Where is (or where will) the Emergency Kit be located?", "Emergency Kit") or ""
    additional_items = get_answer("Add any additional emergency kit items not in the list above (comma-separated):", "Emergency Kit") or ""

    selected_items = st.session_state.get("homeowner_kit_stock", [])
    not_selected_items = st.session_state.get("not_selected_items", [])

    results = st.session_state.get("utility_providers", {})
    electricity_provider_name = results.get("electricity", "")
    natural_gas_provider_name = results.get("natural_gas", "")
    water_provider_name = results.get("water", "")

    flashlights_info = st.session_state.get("flashlights_info", "")
    radio_info = st.session_state.get("radio_info", "")
    food_water_info = st.session_state.get("food_water_info", "")
    important_docs_info = st.session_state.get("important_docs_info", "")
    whistle_info = st.session_state.get("whistle_info", "")
    medications_info = st.session_state.get("medications_info", "")
    mask_info = st.session_state.get("mask_info", "")
    maps_contacts_info = st.session_state.get("maps_contacts_info", "")

    selected_md = "".join(f"- {item}\n" for item in selected_items)
    missing_md = "".join(f"- {item}\n" for item in not_selected_items)
    additional_list = [itm.strip() for itm in additional_items.split(",") if itm.strip()]
    additional_md = "".join(f"- {itm}\n" for itm in additional_list)
    
    kit_summary_line = (
    f"Kit is available at {emergency_kit_location}"
    if emergency_kit_status == "Yes"
    else f"Kit is a work in progress and will be located at {emergency_kit_location}"
    )

    def render_recommended(*items):
        return "".join(f"- {i}\n" for i in items if i and i.strip())

    return f"""
You are an expert assistant generating a city-specific Emergency Preparedness Run Book. First, search the internet for up-to-date local utility providers and their emergency contact information. Then, compose a comprehensive, easy-to-follow guide customized for residents of City: {city}, Zip Code: {zip_code}.

Start by identifying the following utility/service providers for the specified location:
- Internet Provider Name
- Electricity Provider Name
- Natural Gas Provider Name
- Water Provider Name

For each provider, retrieve:
- Company Description
- Customer Service Phone Number
- Customer Service Address (if available)
- Official Website
- Emergency Contact Numbers (specific to outages, leaks, service disruptions)
- Steps to report issues
---

# 🧰 Emergency Kit Summary

## Emergency Location:
{kit_summary_line}

## Kit Inventory:  
{selected_md or "_(none selected)_"}  
## ⚠️ Missing Kit Items (consider adding): 
{missing_md or "_(none missing)_"}  

## Additional User-Added Items: 
{additional_md or "_(none added)_"}  

---

# 📕 Emergency Run Book

## ⚡ 1. Electricity – {electricity_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

### Power Outage Response Guide:
- Steps to follow
- How to report
- Safety precautions
##### Recommended Kit Items:
{render_recommended(flashlights_info, radio_info, food_water_info, important_docs_info)}

---

## 🔥 2. Natural Gas – {natural_gas_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

### Gas Leak Response Guide:
- Signs and precautions
- How to evacuate
- How to report
##### Recommended Kit Items:
{render_recommended(whistle_info, important_docs_info, flashlights_info)}

---

## 💧 3. Water – {water_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

### Water Outage or Leak Guide:
- Detection steps
- Shutoff procedure
#### Recommended Kit Items:
{render_recommended(food_water_info, medications_info, mask_info, important_docs_info)}

---

## 🌐 4. Internet – {internet_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

### Internet Outage Response Guide:
- Troubleshooting
- Reporting
- Staying informed
#### Recommended Kit Items:
{render_recommended(radio_info, maps_contacts_info, important_docs_info)}

Ensure the run book is clearly formatted using Markdown, with bold headers and bullet points. Use ⚠️ to highlight missing kit items.
""".strip()

# Utility to build task metadata
def build_metadata(task_type, label, is_freq=False):
    return {
        "is_task": True,
        "task_label": label,
        "task_type": task_type,
        "frequency_field": is_freq,
        "area": "home"  # or infer from section
    }

def query_utility_providers(test_mode=False): #### Refactored
    """
    Queries Mistral AI for utility providers based on city and ZIP.
    Returns and stores provider names in st.session_state.
    """
    city = get_answer("City", "Home Basics")
    zip_code = get_answer("ZIP Code", "Home Basics")

    if test_mode:
        return {
            "electricity": "Austin Energy",
            "natural_gas": "Atmos Energy",
            "water": "Austin Water"
        }

    prompt = utility_provider_lookup_prompt(city, zip_code)

    try:
        response = client.chat.complete(
            model="mistral-small-latest",
            messages=[UserMessage(content=prompt)],
            max_tokens=1500,
            temperature=0.5,
        )
        content = response.choices[0].message.content
    except Exception as e:
        st.error(f"Error querying Mistral API: {str(e)}")
        content = ""

    results = extract_and_log_providers(content)
    st.session_state["utility_providers"] = results
    return results

def get_home_inputs():
    city = capture_input("City", st.text_input, "Home Basics")
    zip_code = capture_input("ZIP Code", st.text_input, "Home Basics")
    internet_provider = capture_input("Internet Provider", st.text_input, "Home Basics")
    st.session_state.city = city
    st.session_state.zip_code = zip_code
    st.session_state.internet_provider = internet_provider
    return city, zip_code, internet_provider

def get_corrected_providers(results):
    updated = {}

    label_to_key = {
    "Electricity": "electricity",
    "Natural Gas": "natural_gas",
    "Water": "water"
    }

    for label in ["Electricity", "Natural Gas", "Water"]:
        key = label_to_key[label]
        current_value = results.get(key, "") # Use get() to avoid KeyError

        correct_flag = st.checkbox(f"Correct {label} Provider", value=False)
        corrected = st.text_input(
            f"{label} Provider",
            value=current_value,
            disabled=not correct_flag
        )
        if correct_flag and corrected != current_value:
            log_provider_result(label, corrected)
            st.session_state[f"{key}_provider"] = corrected
        updated[key] = corrected if correct_flag else current_value
    return updated

def update_session_state_with_providers(updated):
    st.session_state["utility_providers"] = updated
    for key, value in updated.items():
        st.session_state[f"{key}_provider"] = value

def add_home_schedule_to_docx(doc, schedule_df):
    """
    Adds a grouped home schedule (Source → Date) to DOCX with embedded images inside task cells.

    Args:
        doc (Document): python-docx Document.
        schedule_df (pd.DataFrame): DataFrame with Task, Tag, Date, Source.
    """

    if schedule_df.empty:
        return

    # Set all categories to 'home'
    schedule_df["Category"] = "home"
    schedule_df["Date"] = pd.to_datetime(schedule_df["Date"], errors="coerce")
    schedule_df = schedule_df.sort_values(by=["Source", "Date", "Tag", "Task"])

    # Build image map from session
    image_map = {}
    if "trash_images" in st.session_state:
        for label, img_bytes in st.session_state["trash_images"].items():
            if img_bytes:
                keyword = label.replace(" Image", "").strip().lower()
                image_map[keyword] = img_bytes

    # Begin DOCX layout
    doc.add_page_break()
    doc.add_heading("📆 Home Maintenance Schedule", level=1)

    for source, source_group in schedule_df.groupby("Source"):
        doc.add_heading(f"🗂️ {source}", level=2)

        for date, date_group in source_group.groupby("Date"):
            day = date.strftime("%A")
            date_str = date.strftime("%Y-%m-%d")
            doc.add_heading(f"{day}, {date_str}", level=3)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Task"
            hdr_cells[1].text = "Tag"
            hdr_cells[2].text = "Category"

            for _, row in date_group.iterrows():
                task_text = str(row["Task"])
                tag = str(row["Tag"])
                category = str(row["Category"])
                task_lower = task_text.lower()

                cells = table.add_row().cells

                # Write text first
                paragraph = cells[0].paragraphs[0]
                run = paragraph.add_run(task_text)

                # Match image to task
                for keyword, image_bytes in image_map.items():
                    if keyword in task_lower:
                        try:
                            image_stream = io.BytesIO(image_bytes)
                            image = Image.open(image_stream)
                            image.thumbnail((500, 500))  # Resize if needed

                            resized_stream = io.BytesIO()
                            image.save(resized_stream, format="PNG")
                            resized_stream.seek(0)

                            paragraph.add_run().add_picture(resized_stream, width=Inches(2.5))
                        except Exception as e:
                            paragraph.add_run(f"\n⚠️ Failed to embed image: {e}")
                        break  # One image per task

                cells[1].text = tag
                cells[2].text = category

        doc.add_paragraph("")  # spacing

def generate_convenience_tasks(section_data: dict) -> list[dict]:
    """
    Generates task dictionaries from convenience_seeker section inputs.
    Each dict includes fields required for schedule generation.
    """
    tasks = []
    for service, answers in section_data.items():
        service_name = service.replace(" ", "_").lower()

        company = answers.get(f"{service} Company Name", "")
        phone = answers.get(f"{service} Company Phone Number", "")
        freq = answers.get(f"{service} Frequency", "")
        day = answers.get(f"{service} Day of the Week", "")
        access = answers.get(f"Access Method for {service}", "")
        postproc = answers.get(f"Post-{service} Procedures", "")
        verify = answers.get(f"{service} Crew Identity Verification", "")

        if freq and day:
            task = {
                "Task": f"{service} Service",
                "Category": "home",
                "Area": "Quality-Oriented Services",
                "Source": "convenience_seeker",
                "Tag": service_name,
                "Date": None,  # To be filled in later when expanded to actual calendar dates
                "Day": day,
                "Frequency": freq,
                "Details": {
                    "Company": company,
                    "Phone": phone,
                    "Access": access,
                    "PostProcedure": postproc,
                    "Verification": verify
                }
            }
            tasks.append(task)
    return tasks

def generate_docx_from_split_prompts( ### depreciate 
    prompts: List[str],
    api_key: str,
    *,
    section_titles: Optional[List[str]] = None,
    model: str = "mistral-small-latest",
    doc_heading: str = "Runbook",
    temperature: float = 0.5,
    max_tokens: int = 2048,
    debug: bool = False
) -> Tuple[io.BytesIO, str]:

    if not prompts or not isinstance(prompts, list):
        raise ValueError("🚫 prompts must be a non-empty list of strings.")

    prompts = [p.strip() for p in prompts if p.strip()]
    if not prompts:
        raise ValueError("🚫 All prompts were empty after stripping.")

    combined_output = []

    for i, prompt in enumerate(prompts):
        try:
            if debug:
                st.markdown(f"### 🧾 Prompt Block {i+1}")
                st.code(prompt, language="markdown")

            client = Mistral(api_key=api_key)
            completion = client.chat.complete(
                model=model,
                messages=[SystemMessage(content=prompt)],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            response_text = completion.choices[0].message.content

            if debug:
                st.markdown(f"### 🧾 Raw LLM Response [Block {i + 1}]")
                st.code(response_text, language="markdown")

            title = section_titles[i].strip() if section_titles and i < len(section_titles) else ""
            section_label = f"### {title}" if title else ""
            output_block = f"{section_label}\n\n{response_text.strip()}" if section_label else response_text.strip()
            combined_output.append(output_block)

        except Exception as e:
            st.error(f"❌ Error processing prompt {i + 1}: {e}")
            continue

    full_text = "\n\n".join(combined_output)
    if debug:
        st.markdown("### 🔍 Full Text Passed to DOCX Builder")
        st.code(full_text, language="markdown")

    doc = Document()
    doc.add_heading(doc_heading, 0)
    lines = full_text.splitlines()

    if debug:
        st.markdown("### 🧾 Lines Sent to DOCX")
        st.code("\n".join(lines), language="markdown")

    schedule_sources = get_schedule_placeholder_mapping()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Table placeholders
        if line in schedule_sources:
            df_key = schedule_sources[line]
            schedule_df = st.session_state.get(df_key)

            doc.add_paragraph("")  # spacing
            if isinstance(schedule_df, pd.DataFrame) and not schedule_df.empty:
                if debug:
                    doc.add_paragraph(f"✅ [DEBUG] Inserted table for {line}")
                add_table_from_schedule(doc, schedule_df)
            else:
                doc.add_paragraph("_No schedule available._")
                if debug:
                    st.warning(f"⚠️ No data found for placeholder: {line}")
            doc.add_paragraph("")
            continue

        # Markdown-to-Word formatting
        if line.startswith("##### "):
            doc.add_heading(line[6:].strip(), level=4)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            para = doc.add_paragraph()
            cursor = 0
            for match in re.finditer(r"(\*\*.*?\*\*)", line):
                start, end = match.span()
                if start > cursor:
                    para.add_run(line[cursor:start])
                para.add_run(match.group(1)[2:-2]).bold = True
                cursor = end
            if cursor < len(line):
                para.add_run(line[cursor:])
            para.style.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, full_text

def preview_runbook_output(runbook_text: str, label: str = "📖 Preview Runbook"):
    """
    Shows an expandable markdown preview of the runbook text when a button is clicked.

    Args:
        runbook_text (str): The raw LLM-generated markdown-style text.
        label (str): Button label to trigger the preview.
    """
    # use the following code with home_app_05_23_modified.py
    #if not runbook_text:
    #    st.warning("⚠️ No runbook content available to preview.")
    #    return

    #if st.button(label):
    #    with st.expander("🧠 AI-Generated Runbook Preview", expanded=True):
    #        st.markdown(runbook_text)
    if not runbook_text:
        st.warning("⚠️ No runbook content available to preview.")
        return

    if st.button(label):
        # Optionally add schedule preview
        schedule_df = st.session_state.get("combined_home_schedule_df")
        if isinstance(schedule_df, pd.DataFrame) and not schedule_df.empty:
            schedule_md = add_table_from_schedule_to_markdown(schedule_df)
            runbook_text = runbook_text.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)

        with st.expander("🧠 AI-Generated Runbook Preview", expanded=True):
            st.markdown(runbook_text, unsafe_allow_html=True)

def render_prompt_preview(missing: list, section: str = "home"):
    confirmed = st.session_state.get(f"{section}_user_confirmation", False)

    with st.expander("🧠 AI Prompt Preview (Optional)", expanded=True):
        if missing:
            st.warning(f"⚠️ Cannot generate prompt. Missing: {', '.join(missing)}")
            return

        if not confirmed:
            st.info("☕️ Please check the box to confirm AI prompt generation.")
            return

        prompt = st.session_state.get("generated_prompt", "")
        prompt_blocks = st.session_state.get("prompt_blocks", [])
        schedule_md = st.session_state.get("home_schedule_markdown", "_No schedule available._")

        if not prompt:
            st.warning("⚠️ Prompt not generated yet.")
            return

        # Build combined preview by inserting schedule into the final prompt block
        if prompt_blocks:
            full_preview = "\n\n".join(prompt_blocks)
            full_preview = full_preview.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)
        else:
            full_preview = prompt.replace("<<INSERT_SCHEDULE_TABLE>>", schedule_md)

        # Display the formatted full prompt preview
        st.markdown(full_preview)
        if section == "mail_trash_handling":
            st.markdown("### 📋 Schedule Preview")
            st.markdown(st.session_state.get("home_schedule_markdown", "_No schedule available._"))
        st.success("✅ Prompt ready! This is what will be sent to the LLM.")

def maybe_generate_prompt(section: str) -> Tuple[Optional[str], List[str]]:
    """
    Generate a section-specific prompt and return both the final combined prompt string
    and a list of individual prompt fragments (if any).
    
    Args:
    - section (str): The section of the app (e.g., 'home', 'mail_trash_handling').

    Returns:
    - combined_prompt (str or None)
    - flat_prompts (List[str]) — prompt chunks for use with LLMs
    """
    confirm_key = f"confirm_ai_prompt_{section}"
    confirmed = st.session_state.get(confirm_key, False)

    if st.session_state.get("enable_debug_mode"):
        st.write(f"🧪 maybe_generate_prompt() called for section: `{section}`")
        st.write(f"🧪 Confirmation checkbox state: {confirmed}")

    if not confirmed:
        st.session_state["generated_prompt"] = None
        st.session_state["prompt_blocks"] = []
        return None, []

    # Special case: merge inputs for joint section
    if section == "mail_trash_handling":
        mail_inputs = st.session_state.get("input_data", {}).get("mail", [])
        trash_inputs = st.session_state.get("input_data", {}).get("trash_handling", [])
        merged_inputs = mail_inputs + trash_inputs
        st.session_state["input_data"]["mail_trash_handling"] = merged_inputs
        if st.session_state.get("enable_debug_mode"):
            st.write("📬 [DEBUG] Saved merged input_data['mail_trash_handling']:", merged_inputs)

    # Generate prompt blocks for the specified section
    prompt_blocks = generate_all_prompt_blocks(section)

    # Flatten blocks into a single list
    flat_prompts = []
    for block in prompt_blocks:
        flat_prompts.extend(block if isinstance(block, list) else [block])

    # Join into a single string prompt
    combined_prompt = "\n\n".join(flat_prompts)

    # Save to session state for use elsewhere
    st.session_state["generated_prompt"] = combined_prompt
    st.session_state["prompt_blocks"] = flat_prompts

    return combined_prompt, flat_prompts

def maybe_generate_runbook(section: str, doc_heading: Optional[str] = None):
    """
    Generate a DOCX runbook from prompt blocks for a given section and render download options.

    Parameters:
    - section (str): Logical app section (e.g. 'home', 'mail_trash_handling')
    - doc_heading (Optional[str]): Custom heading for the document
    """
    schedule_placeholder = "<<INSERT_SCHEDULE_TABLE>>"

    # Always generate fresh blocks
    prompt_blocks = generate_all_prompt_blocks(section)
    combined_prompt = "\n\n".join(prompt_blocks)

    # Fallback if prompt blocks are missing but a cached prompt exists
    if not prompt_blocks and st.session_state.get("generated_prompt"):
        prompt_blocks = [st.session_state["generated_prompt"]]
        if st.session_state.get("enable_debug_mode"):
            st.info("⚠️ No prompt blocks returned — using combined prompt from session.")

    # Clean out blank blocks
    prompt_blocks = [b for b in prompt_blocks if b.strip()]
    if not prompt_blocks:
        st.warning(f"⚠️ No valid prompt blocks available for `{section}`. Cannot generate runbook.")
        return

    if doc_heading is None:
        doc_heading = f"{section.replace('_', ' ').title()} Emergency Runbook"

    button_key = f"generate_runbook_button_{section}"
    generate_triggered = st.button("📄 Click Me to Generate Runbook", key=button_key)

    if st.session_state.get("enable_debug_mode"):
        st.write("🔘 Button clicked?", generate_triggered)

    if generate_triggered:
        final_prompt_blocks = prompt_blocks.copy()

        if st.session_state.get("enable_debug_mode"):
            st.markdown("### 🧾 Prompt Blocks Being Sent to LLM")
            for i, block in enumerate(final_prompt_blocks):
                st.code(f"[Block {i + 1}]\n{block}", language="markdown")

        try:
            st.info("⚙️ Calling generate_docx_from_prompt_blocks...")
            buffer, llm_output = generate_docx_from_prompt_blocks(
                blocks=final_prompt_blocks,
                use_llm=True,
                api_key=os.getenv("MISTRAL_TOKEN"),
                doc_heading=doc_heading,
                debug=st.session_state.get("enable_debug_mode", False),
                section=section,  # ✅ Explicit section forwarding (if your version of generate_docx_from_prompt_blocks uses it)
            )
            st.success("✅ DOCX runbook generation completed.")

            if buffer and isinstance(buffer, io.BytesIO) and buffer.getbuffer().nbytes > 0:
                st.session_state[f"{section}_runbook_buffer"] = buffer
                st.session_state[f"{section}_runbook_text"] = llm_output
                st.session_state[f"{section}_runbook_ready"] = True
            else:
                st.warning("⚠️ Runbook buffer is empty. Prompt may contain unresolved placeholders.")
                st.session_state[f"{section}_runbook_ready"] = False

        except Exception as e:
            import traceback
            st.error(f"❌ Exception during runbook generation for `{section}`: {e}")
            st.code(traceback.format_exc(), language="python")
            st.session_state[f"{section}_runbook_ready"] = False

    if st.session_state.get(f"{section}_runbook_ready"):
        st.markdown("___")
        st.write("⏲️ Runbook Ready")
        maybe_render_download(section=section)

def render_schedule_grouped_by_date_then_type_markdown(schedule_df: pd.DataFrame) -> str:
    """
    Returns markdown formatted string grouped by date then task_type with bullets,
    including base64 image links if matched (e.g. trash/recycling).
    """
    if schedule_df.empty:
        return "_No schedule data available._"

    schedule_df["Date"] = pd.to_datetime(schedule_df["Date"], errors="coerce")
    schedule_df = schedule_df.sort_values(by=["Date", "task_type", "Task"])
    schedule_df["Day"] = schedule_df["Date"].dt.strftime("%A")

    lines = []
    current_date = None
    trash_images = st.session_state.get("trash_images", {})

    grouped = schedule_df.groupby(["Date", "task_type"])

    for (date, task_type), group in grouped:
        if date != current_date:
            lines.append(f"### 📅 {date.strftime('%A, %Y-%m-%d')}\n")
            current_date = date

        lines.append(f"#### 📌 {task_type} Schedule\n")

        for _, row in group.iterrows():
            task = row["Task"]
            image_md = ""

            # Attempt to attach image via label match
            for label in ["Outdoor Bin Image", "Recycling Bin Image"]:
                if label.lower().replace(" image", "") in task.lower():
                    uploaded = trash_images.get(label)
                    if uploaded:
                        try:
                            image_bytes = uploaded.getvalue() if hasattr(uploaded, 'getvalue') else uploaded.read()
                            base64_img = base64.b64encode(image_bytes).decode("utf-8")
                            mime = "image/png"
                            image_md = f" ![image](data:{mime};base64,{base64_img})"
                        except Exception:
                            image_md = " ⚠️ (image error)"
                    break

            lines.append(f"- {task}{image_md}")

        lines.append("")  # spacing

    return "\n".join(lines)

def render_schedule_grouped_by_date_then_type(doc: Document, combined_df: pd.DataFrame):
    """
    Adds a single combined schedule to the DOCX, grouped by Date then task_type,
    with image support for matching task labels.
    """
    if combined_df.empty:
        doc.add_paragraph("_No scheduled tasks available._")
        return

    combined_df["Date"] = pd.to_datetime(combined_df["Date"], errors="coerce")
    combined_df = combined_df.sort_values(by=["Date", "task_type", "Task"])
    combined_df["Day"] = combined_df["Date"].dt.strftime("%A")

    trash_images = st.session_state.get("trash_images", {})

    current_date = None
    grouped = combined_df.groupby(["Date", "task_type"])

    for (date, task_type), group_df in grouped:
        if date != current_date:
            if current_date is not None:
                doc.add_paragraph("")  # spacing b_

def autolog_location_inputs(city, zip_code):
    """
    Optionally logs city and zip code inputs to 'Home Basics' section if not already logged.
    """
    now = datetime.now().isoformat()

    # Ensure input_data exists
    if "input_data" not in st.session_state:
        st.session_state["input_data"] = {}

    # Ensure the 'Home Basics' section exists
    if "Home Basics" not in st.session_state["input_data"]:
        st.session_state["input_data"]["Home Basics"] = []

    def log_if_missing(label, value):
        if value and not any(entry["question"] == label for entry in st.session_state["input_data"]["Home Basics"]):
            # Add to input_data
            st.session_state["input_data"]["Home Basics"].append({
                "question": label,
                "answer": value,
                "timestamp": now
            })

            # Also add to interaction_log if enabled
            st.session_state.setdefault("interaction_log", []).append({
                "timestamp": now,
                "user_id": st.session_state.get("user_id", "anonymous"),
                "session_id": st.session_state.get("session_id", "anonymous"),
                "action": "autolog",
                "question": label,
                "answer": value,
                "section": "Home Basics"
            })

    log_if_missing("City", city)
    log_if_missing("ZIP Code", zip_code)

