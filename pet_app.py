import streamlit as st
from mistralai import Mistral, UserMessage, SystemMessage
import csv
import io
from datetime import datetime, timedelta
from docx import Document
from collections import defaultdict
import json
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import pandas as pd
import re
from utils.data_access import (
    get_saved_pets_by_species,
)

# Set your MISTRAL API key

api_key = os.getenv("MISTRAL_TOKEN")
client = Mistral(api_key=api_key)

if not api_key:
    api_key = st.text_input("Enter your Mistral API key:", type="password")

if api_key:
    st.success("API key successfully loaded.")
else:
   st.error("API key is not set.")

# --------------------------- #
# UTILITIES
# --------------------------- #

def check_bingo(answers):
    for i in range(7):
        if all(answers[i][j] for j in range(7)) or all(answers[j][i] for j in range(7)):
            return True
    if all(answers[i][i] for i in range(7)) or all(answers[i][6 - i] for i in range(7)):
        return True
    return False

def flatten_answers_to_dict(questions, answers):
    return {
        questions[row * 7 + col]['label']: answers[row][col]
        for row in range(7) for col in range(7)
        if answers[row][col].strip()
    }

def get_pet_name(answers):
    try:
        return answers[0][0].strip() or "UnnamedPet"
    except:
        return "UnnamedPet"

def convert_to_csv(data_list):
    if not data_list:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=data_list[0].keys())
    writer.writeheader()
    writer.writerows(data_list)
    return output.getvalue()

def export_all_pets_to_docx(saved_pets, species):
    doc = Document()
    doc.add_heading(f"{species.capitalize()} Care Report", 0)
    for i, pet in enumerate(saved_pets, 1):
        name = pet.get("🐕 Pet Name", f"{species.capitalize()} #{i}")
        doc.add_heading(f"{i}. {name}", level=1)
        categories = {}
        for question, answer in pet.items():
            for q in all_question_metadata[species]:
                if q["label"] == question:
                    cat = q["category"]
                    categories.setdefault(cat, []).append((q["label"], answer))
                    break
        for cat, qas in categories.items():
            doc.add_heading(cat, level=2)
            for label, answer in qas:
                doc.add_paragraph(f"{label}: {answer}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def load_pet_for_edit(index, state_prefix, questions):
    selected = st.session_state[f"saved_{state_prefix}s"][index]
    new_answers = [['' for _ in range(7)] for _ in range(7)]
    for i in range(49):
        row, col = divmod(i, 7)
        label = questions[i]['label']
        new_answers[row][col] = selected.get(label, "")
    st.session_state[f"answers_{state_prefix}"] = new_answers
    st.session_state[f"editing_index_{state_prefix}"] = index
    st.experimental_rerun()

def render_pet_tab(pet_type, questions, state_prefix):
    st.header(f"{pet_type.capitalize()} Care Form")

    st.session_state.setdefault(f'answers_{state_prefix}', [['' for _ in range(7)] for _ in range(7)])
    st.session_state.setdefault(f'saved_{state_prefix}s', [])
    st.session_state.setdefault(f'{state_prefix}_index', 1)
    st.session_state.setdefault(f'editing_index_{state_prefix}', None)

    answers = st.session_state[f'answers_{state_prefix}']
    saved_pets = st.session_state[f'saved_{state_prefix}s']
    editing_index = st.session_state[f'editing_index_{state_prefix}']
    index = st.session_state[f'{state_prefix}_index']

    st.subheader(f"{'✏️ Editing' if editing_index is not None else '🐾'} {pet_type.capitalize()} #{index}")

    bingo_board = [questions[i:i + 7] for i in range(0, 49, 7)]

    for row_index in range(7):
        cols = st.columns(7)
        for col_index in range(7):
            question = bingo_board[row_index][col_index]
            q_label = question['label']
            current_value = answers[row_index][col_index]
            with cols[col_index]:
                with st.expander(q_label):
                    new_value = st.text_area(
                        "Answer Here",
                        key=f"{state_prefix}_q{col_index}_{row_index}",
                        value=current_value,
                        placeholder="Enter your answer",
                        label_visibility="collapsed"
                    )
                    answers[row_index][col_index] = new_value
                    st.markdown("✔️ Answered" if new_value else "❓ Not Answered")

    if check_bingo(answers):
        st.success("🎉 Bingo complete!")

        if st.button("💾 Save Entry", key=f"save_{state_prefix}"):
            data = flatten_answers_to_dict(questions, answers)
            if editing_index is not None:
                st.session_state[f"saved_{state_prefix}s"][editing_index] = data
                st.session_state[f"editing_index_{state_prefix}"] = None
                st.success("✅ Entry updated!")
            else:
                st.session_state[f"saved_{state_prefix}s"].append(data)
                st.session_state[f"{state_prefix}_index"] += 1
                st.success("✅ New entry saved!")
            st.session_state[f"answers_{state_prefix}"] = [['' for _ in range(7)] for _ in range(7)]
            st.experimental_rerun()

        pet_name = get_pet_name(answers)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        filename = f"{pet_name}_{timestamp}.csv"
        csv_data = convert_to_csv([flatten_answers_to_dict(questions, answers)])
        st.download_button("⬇️ Download This Entry as CSV", csv_data, file_name=filename, mime="text/csv", key=f"dl_single_{state_prefix}")

    if saved_pets:
        st.markdown("### 📋 Saved Entries:")
        for i, pet in enumerate(saved_pets):
            name = pet.get("🐕 Pet Name", f"{pet_type.capitalize()} #{i+1}")
            cols = st.columns([5, 1])
            cols[0].markdown(f"**{i+1}. {name}**")
            if cols[1].button("✏️ Edit", key=f"edit_{state_prefix}_{i}"):
                load_pet_for_edit(i, state_prefix, questions)

        all_csv = convert_to_csv(saved_pets)
        st.download_button(f"⬇️ Download All {pet_type.capitalize()}s as CSV", all_csv, file_name=f"all_{pet_type}s.csv", mime="text/csv", key=f"dl_all_csv_{state_prefix}")
        docx_buf = export_all_pets_to_docx(saved_pets, pet_type)
        st.download_button(f"📄 Download All {pet_type.capitalize()}s as DOCX", docx_buf, file_name=f"all_{pet_type}s.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_all_docx_{state_prefix}")

def select_runbook_date_range():
    """
    Display a section for the user to choose a date range or timeframe for runbook generation.

    Returns:
        tuple: (choice, start_date, end_date)
            - choice: selected option from the radio
            - start_date: calculated or selected start date
            - end_date: calculated or selected end date
    """
    st.subheader("Choose Date(s) or Timeframe")
    st.write("Choose a timeframe you would like a runbook generated for.")

    options = ["Pick Dates", "Weekdays Only", "Weekend Only", "Default"]
    choice = st.radio("Choose an option:", options)

    start_date, end_date = None, None
    today = datetime.now().date()

    if choice == "Pick Dates":
        start_date = st.date_input("Select Start Date:", today)
        end_date = st.date_input("Select End Date:", today + timedelta(days=7))
        st.write(f"📅 You selected specific dates from **{start_date}** to **{end_date}**.")

    elif choice == "Weekdays Only":
        # Find the next Monday
        days_ahead = (0 - today.weekday() + 7) % 7  # 0 = Monday
        start_date = today + timedelta(days=days_ahead)
        end_date = start_date + timedelta(days=4)
        st.info(f"📅 Auto-selected **Weekdays Only**: {start_date} to {end_date}")

    elif choice == "Weekend Only":
        # Find the next Saturday
        days_ahead = (5 - today.weekday() + 7) % 7  # 5 = Saturday
        start_date = today + timedelta(days=days_ahead)
        end_date = start_date + timedelta(days=1)
        st.info(f"📅 Auto-selected **Weekend Only**: {start_date} to {end_date}")

    elif choice == "Default":
        start_date = today
        end_date = today + timedelta(days=7)
        st.info(f"📅 Default schedule: {start_date} to {end_date}")

    else:
        st.warning("Invalid choice.")

    return choice, start_date, end_date

def extract_pet_scheduled_tasks(questions, saved_pets, start_date, end_date):
    """
    Extracts daily, weekly, and one-time pet care tasks.
    Returns a long-form DataFrame with Day, Pet, Task, Tag.
    """
    schedule_rows = []

    # Recurring task keyword categories
    recurring_keywords = {
        "feeding": ["feed", "feeding schedule", "food", "treats"],
        "medication": ["medication", "pill", "dose", "schedule"],
        "walk": ["walk", "walking", "exercise"],
        "grooming": ["brush", "bathing", "grooming", "nail", "teeth", "ear cleaning"],
        "play": ["play", "toys", "activities"],
        "litter": ["litter", "waste", "box"],
        "water": ["water", "bowl"]
    }

    # Keywords to flag one-time events with dates
    one_time_keywords = ["pill", "appointment", "check-up", "vaccination", "vet", "visit"]

    def label_matches(label, keywords):
        return any(kw in label.lower() for kw in keywords)

    def extract_dates(text):
        # Parse dates like "June 5", "6/5/24", etc.
        patterns = [
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}(?:,\s*\d{4})?",
            r"\b\d{1,2}/\d{1,2}(?:/\d{2,4})?"
        ]
        matches = []
        for pattern in patterns:
            matches += re.findall(pattern, text, flags=re.IGNORECASE)
        return matches

    def extract_weekday_mentions(text):
        weekdays = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        return [day for day in weekdays if day.lower() in text.lower()]

    def normalize_date(date_str):
        try:
            return pd.to_datetime(date_str, errors="coerce").date()
        except:
            return None

    emoji_tags = {
        "[Daily]": "🔁 Daily",
        "[Weekly]": "📅 Weekly",
        "[One-Time]": "🗓️ One-Time",
        "[Assumed-Daily]": "🔄 Assumed Daily"
    }

    for pet in saved_pets:
        pet_name = pet.get("🐕 Pet Name", "Unnamed Pet")
        for q in questions:
            label = q["label"]
            answer = pet.get(label, "").strip()
            if not answer:
                continue

            base_task = f"{pet_name}: {label} – {answer}"

            # 1. Daily (recurring)
            for keywords in recurring_keywords.values():
                if label_matches(label, keywords):
                    current_date = start_date
                    while current_date <= end_date:
                        schedule_rows.append({
                            "Day": current_date.strftime('%A'),
                            "Pet": pet_name,
                            "Task": f"{label} – {answer}",
                            "Tag": emoji_tags["[Daily]"]
                        })
                        current_date += timedelta(days=1)
                    break

            # 2. Weekly (e.g., "every Monday")
            for weekday in extract_weekday_mentions(answer):
                schedule_rows.append({
                    "Day": weekday,
                    "Pet": pet_name,
                    "Task": f"{label} – {answer}",
                    "Tag": emoji_tags["[Weekly]"]
                })

            # 3. One-Time (e.g., "pill on June 5")
            if label_matches(label, one_time_keywords):
                for ds in extract_dates(answer):
                    parsed_date = normalize_date(ds)
                    if parsed_date and start_date <= parsed_date <= end_date:
                        schedule_rows.append({
                            "Day": parsed_date.strftime('%A'),
                            "Pet": pet_name,
                            "Task": f"{label} – {answer}",
                            "Tag": emoji_tags["[One-Time]"]
                        })

    df = pd.DataFrame(schedule_rows)
    df = df.sort_values(by=["Day", "Tag", "Pet"]).reset_index(drop=True)
    return df

# --------------------------- #
# QUESTION SETS
# --------------------------- #

dog_questions = [
    {"id": "name", "label": "🐕 Dog's Name", "category": "Basic Info"},
    {"id": "vet_contact", "label": "🏥 Vet Contact Info (Name, Phone Number, Address)", "category": "Health"},
    {"id": "food_type", "label": "🥣 Food brand/type your dog eats", "category": "Feeding"},
    {"id": "walk_routine", "label": "🧳 Walk Routine (Time, Duration, Location, Behavior)", "category": "Exercise"},
    {"id": "bathing_schedule", "label": "🛁 Bathing Schedule", "category": "Grooming"},
    {"id": "favorite_toys", "label": "🧸 Favorite Toys", "category": "Behavior"},
    {"id": "training_goals", "label": "🎯 Training Goals", "category": "Training"},
    {"id": "breed", "label": "🦴 Breed", "category": "Basic Info"},
    {"id": "emergency_vet", "label": "⛑️ Emergency Vet Contact Info (Name, Phone Number, Address)", "category": "Health"},
    {"id": "meal_portion", "label": "🍖 Meal portion size", "category": "Feeding"},
    {"id": "walk_location", "label": "📍 Favorite Walk Location", "category": "Exercise"},
    {"id": "brushing_schedule", "label": "💈 Brushing Schedule", "category": "Grooming"},
    {"id": "play_styles", "label": "🐶 Play Styles", "category": "Behavior"},
    {"id": "training_challenges", "label": "🥁 Training Challenges", "category": "Training"},
    {"id": "age_weight", "label": "🎂 Age and Weight", "category": "Basic Info"},
    {"id": "medical_conditions", "label": "💊 Medical Conditions and Allergies", "category": "Health"},
    {"id": "feeding_schedule", "label": "🕥 Feeding Schedule", "category": "Feeding"},
    {"id": "walking_equipment", "label": "🐶 Walking Equipment", "category": "Exercise"},
    {"id": "nail_trimming", "label": "💅 Nail Trimming", "category": "Grooming"},
    {"id": "favorite_activities", "label": "🎾 Favorite Activities", "category": "Behavior"},
    {"id": "training_methods", "label": "📚 Training Methods", "category": "Training"},
    {"id": "microchip", "label": "🔖 Microchip Number", "category": "Basic Info"},
    {"id": "medication_schedule", "label": "🕥 Medication Dosage and Schedule", "category": "Health"},
    {"id": "treats", "label": "🍗 Treats/Snacks", "category": "Feeding"},
    {"id": "walk_behavior", "label": "🐾 Walk Behavior", "category": "Exercise"},
    {"id": "ear_cleaning", "label": "👂 Ear Cleaning", "category": "Grooming"},
    {"id": "fear_triggers", "label": "❗ Fear Triggers", "category": "Behavior"},
    {"id": "trainer_contact", "label": "🏫 Trainer Contact (Name, Phone, Email)", "category": "Training"},
    {"id": "appearance", "label": "🖼️ Appearance Description", "category": "Basic Info"},
    {"id": "medication_instructions", "label": "💊 Medication Instructions", "category": "Health"},
    {"id": "treat_frequency", "label": "🕥 Treat Frequency", "category": "Feeding"},
    {"id": "walk_treats", "label": "🍭 Treats for Walks", "category": "Feeding"},
    {"id": "teeth_brushing", "label": "🦷 Teeth Brushing", "category": "Grooming"},
    {"id": "commands_known", "label": "📢 Commands Known", "category": "Training"},
    {"id": "travel_setup", "label": "🌴 Travel Setup", "category": "Logistics"},
    {"id": "spay_neuter", "label": "✂️ Spayed/Neutered", "category": "Basic Info"},
    {"id": "health_history", "label": "🗄️ Health History", "category": "Health"},
    {"id": "water_refill", "label": "💧 Water Refill Schedule", "category": "Feeding"},
    {"id": "sleep_schedule", "label": "💤 Sleep Schedule", "category": "Routine"},
    {"id": "special_grooming", "label": "🌟 Special Grooming Needs", "category": "Grooming"},
    {"id": "behavior_issues", "label": "🔍 Behavioral Issues", "category": "Behavior"},
    {"id": "car_sickness", "label": "🚗 Car Sickness?", "category": "Health"},
    {"id": "adoption_info", "label": "🏘️ Adoption Info (Place, Date)", "category": "Basic Info"},
    {"id": "next_checkup", "label": "📆 Next Check-up Date", "category": "Health"},
    {"id": "sitter_instructions", "label": "📋 Sitter Instructions", "category": "Logistics"},
    {"id": "special_play", "label": "🎾 Special Playtimes", "category": "Behavior"},
    {"id": "walker_contact", "label": "🚶‍♂️ Walker Contact (Name, Phone, Email)", "category": "Logistics"},
    {"id": "socialization", "label": "🐶 Socialization", "category": "Behavior"},
    {"id": "sitter_contact", "label": "🏠 Sitter Contact (Name, Phone, Email)", "category": "Logistics"},
]

cat_questions = [
    {"id": "name", "label": "🐈 Cat's Name", "category": "Basic Info"},
    {"id": "vet_contact", "label": "🏥 Vet Contact Info (Name, Phone Number, Address)", "category": "Health"},
    {"id": "food_type", "label": "🥣 Describe the brand/type of food your cat eats", "category": "Feeding"},
    {"id": "scratching_posts", "label": "🧳 Environment Enrichment  (Scratching Posts/Pads)", "category": "Enrichment"},
    {"id": "bathing_schedule", "label": "🛁 Bathing Schedule", "category": "Grooming"},
    {"id": "favorite_toys", "label": "🧸 Environment Enrichment (Favorite Toys)", "category": "Enrichment"},
    {"id": "training_goals", "label": "🎯 Current Training Goals", "category": "Training"},
    {"id": "breed", "label": "🐱 Name the Breed/Type", "category": "Basic Info"},
    {"id": "emergency_vet", "label": "⛑️ Emergency Vet Contact Info (Name, Phone Number, Address)", "category": "Health"},
    {"id": "meal_portion", "label": "🥘 Describe the portion size for each meal", "category": "Feeding"},
    {"id": "outdoor_access", "label": "📍 Environment Enrichment (Outdoor Access - Yes/No, Supervised/Unsupervised)", "category": "Enrichment"},
    {"id": "brushing_schedule", "label": "💈 Brushing Schedule", "category": "Grooming"},
    {"id": "perches_1", "label": "😸 Environment Enrichment (Cat Tree/Perches)", "category": "Enrichment"},
    {"id": "training_challenges", "label": "🥁 Training Progress/Challenges", "category": "Training"},
    {"id": "age_weight", "label": "🎂 Cat’s Age and Weight", "category": "Basic Info"},
    {"id": "medical_conditions", "label": "💊 List all medical conditions or allergies", "category": "Health"},
    {"id": "feeding_schedule", "label": "🕥 Feeding Schedule", "category": "Feeding"},
    {"id": "perches_2", "label": "🐱 Environment Enrichment (Cat Tree/Perches)", "category": "Enrichment"},
    {"id": "nail_trimming", "label": "💅 Nail Trimming", "category": "Grooming"},
    {"id": "favorite_activities", "label": "🎾 Favorite Activities", "category": "Behavior"},
    {"id": "litter_cleaning", "label": "📚 Waste Management (Litter Box Cleaning Routine, Waste Disposal Method)", "category": "Litter & Hygiene"},
    {"id": "microchip", "label": "🔖 Cat’s microchip number", "category": "Basic Info"},
    {"id": "medication_schedule", "label": "🕥 Medication Schedule with Dosage", "category": "Health"},
    {"id": "treats", "label": "🍭 Name your Cat’s treats or snacks", "category": "Feeding"},
    {"id": "litter_box", "label": "🐾 Litter Box (Type, Brand/Type, Location)", "category": "Litter & Hygiene"},
    {"id": "ear_cleaning", "label": "👂 Ear Cleaning", "category": "Grooming"},
    {"id": "fear_triggers", "label": "❗ Fear/Anxiety Triggers", "category": "Behavior"},
    {"id": "placeholder_1", "label": "🏫 Placeholder Question 1", "category": "Other"},
    {"id": "appearance", "label": "🖼️ Describe the Cat’s Appearance from Memory", "category": "Basic Info"},
    {"id": "medication_instructions", "label": "💊 Medication Delivery Instructions", "category": "Health"},
    {"id": "placeholder_2", "label": "🕥 Placeholder Question 2", "category": "Other"},
    {"id": "treat_frequency", "label": "🍭 When are treats or snacks given?", "category": "Feeding"},
    {"id": "teeth_brushing", "label": "🦷 Teeth Brushing", "category": "Grooming"},
    {"id": "commands_known", "label": "📢 Commands Known", "category": "Training"},
    {"id": "travel_setup", "label": "🌴 Travel carte or car travel setup", "category": "Logistics"},
    {"id": "spay_neuter", "label": "✂️ Cat is Spayed or Neutered", "category": "Health"},
    {"id": "health_history", "label": "🗄️ Health & Vaccination History", "category": "Health"},
    {"id": "water_bowl", "label": "💧 Water bowl refill schedule", "category": "Feeding"},
    {"id": "sleep_schedule", "label": "💤 Sleep Schedule", "category": "Routine"},
    {"id": "special_grooming", "label": "🌟 Special Grooming Needs", "category": "Grooming"},
    {"id": "behavior_issues", "label": "🔍 Behavioral Issues", "category": "Behavior"},
    {"id": "car_sickness", "label": "🚗 Car Sickness?", "category": "Health"},
    {"id": "adoption_info", "label": "🏘️ Place and date the Cat was adopted", "category": "Basic Info"},
    {"id": "next_checkup", "label": "📆 Date of Cat’s next check-up or vaccination", "category": "Health"},
    {"id": "sitter_instructions", "label": "Bonus: Special Instructions for Sitters", "category": "Logistics"},
    {"id": "special_playtimes", "label": "🎾 Special Activities or Playtimes", "category": "Behavior"},
    {"id": "groomer_contact", "label": "🚶‍♂️ Bonus: Pet Groomer Contact Info", "category": "Logistics"},
    {"id": "socialization", "label": "🐱 Socialization with other animals, children, and strangers", "category": "Behavior"},
    {"id": "sitter_contact", "label": "🏠 Bonus: Pet Sitter Contact Info", "category": "Logistics"}
]

all_question_metadata = {
    "dog": dog_questions,
    "cat": cat_questions
}

# --------------------------- #
# PREVIEW --- need to add start_date, end_date, questions, answers
# --------------------------- #

def generate_prompt_for_all_pets_combined(saved_data_by_species, metadata_by_species, start_date, end_date,schedule_markdown=None):
    """
    Generate a multi-pet, multi-species AI prompt, grouped by species and pet,
    with optional appended markdown-formatted schedule.

    Args:
        saved_data_by_species (dict): e.g. { "dog": [...], "cat": [...] }
        metadata_by_species (dict): e.g. { "dog": [...], "cat": [...] }
        start_date (date): care period start
        end_date (date): care period end
        schedule_markdown (str): optional markdown string from schedule_df.to_markdown()

    Returns:
        str: Final prompt for AI input
    """
    prompt_sections = []

    all_saved_pets = []
    all_questions = []

    for species in ["dog", "cat"]:
        pets = saved_data_by_species.get(species, [])
        if not pets:
            continue

        metadata = metadata_by_species[species]
        all_saved_pets.extend(pets)
        all_questions.extend(metadata)

        species_icon = "🐶" if species == "dog" else "🐱"
        species_label = "Dogs" if species == "dog" else "Cats"

        prompt_sections.append(f"# {species_icon} {species_label}")

        for i, pet in enumerate(pets, 1):
            name = pet.get("🐕 Pet Name", f"{species_label[:-1]} #{i}")
            pet_block = [f"## {species_icon} {name}"]
            categories = defaultdict(list)

            for question, answer in pet.items():
                for q in metadata:
                    if q["label"] == question:
                        cat = q.get("category", "Additional Notes")
                        categories[cat].append(f"**{q['label']}**: {answer}")
                        break

            for category, qa_list in categories.items():
                pet_block.append(f"### {category}")
                pet_block.extend(qa_list)
                pet_block.append("")  # spacing

            prompt_sections.extend(pet_block)

    # Compile markdown sections
    prompt_body = "\n".join(prompt_sections)

    # Base prompt
    final_prompt = f"""
You are an AI assistant tasked with generating a **comprehensive Multi-Pet Sitting Runbook** for both dogs and cats during the care period from **{start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}**.

---

### 📝 Instructions:
- Group pets by species (dogs first, then cats).
- For each pet, create a detailed care plan organized by category.
- Use markdown-style formatting with headings and subheadings.
- Be clear, friendly, and actionable. Use bullet points for routines.
- If any section is missing, insert _"No information provided."_.

---

### 📦 Structured Pet Input:

{prompt_body}
"""

    # Append schedule markdown if available
    if schedule_markdown:
        final_prompt += f"""

---

### 📆 Weekly Pet Care Schedule (Day × Task Grid)

{schedule_markdown}
"""

    # Final instruction to LLM
    final_prompt += "\n\nNow generate the full runbook grouped by species and pet."

    return final_prompt
    
def generate_docx_from_prompt(prompt: str, api_key: str, doc_heading: str = "Pet Sitting Runbook") -> tuple[io.BytesIO, str]:
    """
    Generate a styled DOCX runbook from a full prompt string using the Mistral API.

    Args:
        prompt: Full LLM prompt (already assembled).
        api_key: Mistral API token.
        doc_heading: Title of the DOCX document.

    Returns:
        (BytesIO stream, raw AI output text)
    """
    try:
        with st.spinner("🧠 Generating runbook with Mistral..."):
            client = Mistral(api_key=api_key)
            completion = client.chat.complete(
                model="mistral-small-latest",
                messages=[
                    ChatMessage(role=Role.SYSTEM, content="You are a helpful assistant that formats pet sitting instructions into clear, organized prose."),
                    ChatMessage(role=Role.USER, content=prompt)
                ],
                max_tokens=3000,
                temperature=0.5,
            )
            runbook_text = completion.choices[0].message.content
    except Exception as e:
        st.error(f"Mistral API error: {e}")
        return None, ""

    # Convert markdown-style output into a .docx
    doc = Document()
    doc.add_heading(doc_heading, 0)

    for line in runbook_text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[:2].isdigit() and line[2:4] == ". ":
            doc.add_paragraph(line[4:].strip(), style="List Number")
        else:
            para = doc.add_paragraph(line)
            para.style.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, runbook_text

# --------------------------- #
# APP LAYOUT
# --------------------------- #

st.title("🐾 Pet Care Bingo Game")

with st.expander("📖 How to Use This Game", expanded=False):
    st.markdown("""
    Use this form to collect detailed care information for your dogs and cats.

    ### How It Works:
    - Fill out the bingo grid for one pet.
    - When you complete a row/column/diagonal, you can save and export the info.
    - Switch between tabs to manage dogs and cats separately.
    - You can edit saved entries, and download all data as CSV or DOCX.

    Tip: Use this app to prepare pet care docs for sitters, walkers, or emergency planning!
    """)

tab1, tab2, tab3 = st.tabs(["🐶 Dogs", "🐱 Cats", "Preview"])

with tab1:
    render_pet_tab("dog", dog_questions, "dog")

with tab2:
    render_pet_tab("cat", cat_questions, "cat")

with tab3:
    st.header("An owl")
    st.image("https://static.streamlit.io/examples/owl.jpg", width=200)

    ##### work in progress ######

    # Section for date range selection
choice, start_date, end_date = select_runbook_date_range()

# Dynamically gather pet data from session state
saved_data_by_species = get_saved_pets_by_species()
metadata_by_species = {
    species: st.session_state.get(f"{species}_questions", [])
    for species in saved_data_by_species.keys()
}

for species in ["dog", "cat"]:
    saved_key = f"saved_{species}s"
    question_key = f"{species}_questions"

    if saved_key in st.session_state and st.session_state[saved_key]:
        saved_data_by_species[species] = st.session_state[saved_key]
        metadata_by_species[species] = st.session_state[question_key]  # This should be loaded earlier

# ✅ Step 2: Confirm data exists before proceeding
if saved_data_by_species:
    st.markdown("## 🧠 Generate Combined Multi-Pet Runbook")

    # ✅ Step 3: Sanity Check Before Scheduling or Prompting
    required_keys = ["dog_questions", "cat_questions", "saved_dogs", "saved_cats"]
    missing_keys = [k for k in required_keys if k not in st.session_state]

    if missing_keys:
        st.error(f"⚠️ Missing session keys: {', '.join(missing_keys)}. Please complete all forms before generating the runbook.")
    elif not st.session_state["saved_dogs"] and not st.session_state["saved_cats"]:
        st.warning("🐾 No pets have been saved yet. Please fill out at least one form.")
    elif "start_date" not in st.session_state or "end_date" not in st.session_state:
        st.warning("📅 Please select a care date range before generating the runbook.")
    else:
        # ✅ All checks passed — safe to build schedule
        schedule_df, warnings = extract_pet_scheduled_tasks(
            questions=sum(metadata_by_species.values(), []),  # flatten metadata
            saved_pets=sum(saved_data_by_species.values(), []),  # flatten pets
            valid_dates=valid_dates     
            )       
        
        if warnings:
            for note in warnings:
                st.warning(note)

    prompt = generate_prompt_for_all_pets_combined(
        saved_data_by_species=saved_data_by_species,
        metadata_by_species=metadata_by_species,
        start_date=start_date,
        end_date=end_date,
        schedule_markdown=schedule_df.to_markdown(index=False)
        )

    with st.expander("📋 Review AI Prompt", expanded=True):
        st.code(prompt, language="markdown")

    # Ask user to confirm prompt
    confirm = st.checkbox("✅ I confirm this AI prompt is correct and ready for generation.")

    if confirm:
        if st.button("🚀 Generate Pet Sitting Runbook"):
            docx_buffer, runbook_text = generate_docx_from_prompt(
                prompt=prompt,
                api_key=os.getenv("MISTRAL_TOKEN"),
                doc_heading="Multi-Pet Sitting Runbook"  # ✅ added missing comma
            )

            if docx_buffer:
                st.subheader("📋 Runbook Preview")
                st.code(runbook_text, language="markdown")
                st.download_button(
                    label="📄 Download Runbook (.docx)",
                    data=docx_buffer,
                    file_name="multi_pet_runbook.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
else:
    st.info("🪄 Fill out at least one dog or cat form to enable combined runbook generation.")

