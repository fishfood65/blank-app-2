import pytest
from unittest.mock import patch, MagicMock
from utils.runbook_generator_helpers import (
    preview_runbook_output, # test included
    generate_docx_from_split_prompts, # test included
    generate_docx_from_text, # test included
    maybe_generate_prompt, # test included
    render_prompt_preview, # test included
    maybe_render_download,
    maybe_generate_runbook # test included
    )
from utils.common_helpers import get_schedule_utils
from docx import Document
import tempfile
import types
import pandas as pd
from io import BytesIO
import os
import io
from contextlib import contextmanager
from utils import runbook_generator_helpers as rh

def test_preview_runbook_output_does_not_crash(monkeypatch):
    import streamlit as st
    monkeypatch.setattr(st, "button", lambda *args, **kwargs: False)
    monkeypatch.setattr(st, "expander", lambda label, expanded=True: None)
    monkeypatch.setattr(st, "markdown", lambda x: x)
    monkeypatch.setattr(st, "warning", lambda x: x)
    
    # Just verify it doesn't raise
    preview_runbook_output("### Heading\n- Item")

class FakeChat:
    def complete(self, model, messages, max_tokens, temperature):
        print("✅ FakeChat.complete called with message:", messages[0].content)  # Debug
        class Message:
            def __init__(self):
                self.content = "Mocked response content"
        class Choice:
            def __init__(self):
                self.message = Message()
        class Response:
            def __init__(self):
                self.choices = [Choice()]
        return Response()

class FakeClient:
    def __init__(self, api_key):
        print("✅ FakeClient created with api_key:", api_key)  # Debug
        self.api_key = api_key
        self.chat = FakeChat()

@contextmanager
def fake_spinner(msg):
    print("SPINNER:", msg)
    yield

def test_generate_docx_from_split_prompts(monkeypatch):
    import streamlit as st
    import utils.runbook_generator_helpers as rh

    # Fix context manager error
    monkeypatch.setattr(st, "spinner", fake_spinner)
    monkeypatch.setattr(st, "error", lambda *args, **kwargs: print("STREAMLIT ERROR:", args[0]))
    monkeypatch.setattr("utils.runbook_generator_helpers.Mistral", lambda api_key: FakeClient(api_key))

    prompts = ["This is prompt 1.", "This is prompt 2."]
    section_titles = ["Section 1", "Section 2"]
    api_key = "fake-key"

    buffer, full_text = rh.generate_docx_from_split_prompts(
        prompts,
        api_key,
        section_titles=section_titles,
        doc_heading="Test Runbook"
    )

    print("📄 FULL TEXT RETURNED:", full_text)
    assert isinstance(buffer, io.BytesIO)
    assert "Mocked response content" in full_text

def test_get_schedule_utils_keys_and_types():
    utils = get_schedule_utils()
    expected_keys = [
        "emoji_tags",
        "weekday_to_int",
        "extract_week_interval",
        "extract_weekday_mentions",
        "extract_dates",
        "normalize_date",
        "determine_frequency_tag"
    ]
    for key in expected_keys:
        assert key in utils
    assert callable(utils["extract_week_interval"])
    assert isinstance(utils["emoji_tags"], dict)

def test_generate_docx_from_text(monkeypatch: pytest.MonkeyPatch):
    # Dummy home schedule DataFrame to test schedule table insertion
    sample_df = pd.DataFrame([
        {"Date": "2024-01-01", "Task": "Test task", "Tag": "test", "Category": "home", "Source": "unit test"}
    ])

    # Create a fake streamlit module with session_state
    fake_st = types.SimpleNamespace()
    fake_st.session_state = {
        "home_schedule_df": sample_df
    }

    # Monkeypatch streamlit used inside generate_docx_from_text
    import utils.common_helpers as uh
    monkeypatch.setattr(rh, "st", fake_st)

    # Text input with the insertion marker
    test_text = "<<INSERT_SCHEDULE_TABLE>>"
    buffer = rh.generate_docx_from_text(test_text, doc_heading="Test Runbook")

    # Read resulting DOCX
    doc = Document(buffer)

    # Confirm heading exists
    headings = [p.text for p in doc.paragraphs if p.style.name in ("Title", "Heading 1", "Heading 2")]
    assert "Test Runbook" in headings, "❌ Missing heading in DOCX"

    # Confirm the table was inserted
    assert len(doc.tables) == 1, "❌ Expected 1 table inserted from schedule"

def test_maybe_generate_prompt_for_home_section():
    fake_st = types.SimpleNamespace()
    fake_st.session_state = {
        "confirm_ai_prompt_home": True
    }

    with patch("utils.runbook_generator_helpers.st", fake_st), \
         patch("utils.runbook_generator_helpers.utilities_emergency_runbook_prompt", return_value="Prompt for home"):

        combined, prompts = maybe_generate_prompt("home")

        assert combined == "Prompt for home"
        assert prompts == ["Prompt for home"]
        assert fake_st.session_state["generated_prompt"] == "Prompt for home"

def test_render_prompt_preview_confirmed_with_prompt(monkeypatch):
    # Prepare fake Streamlit module
    fake_st = MagicMock()
    fake_st.session_state = {
        "home_user_confirmation": True,
        "generated_prompt": "Test prompt"
    }

    monkeypatch.setitem(fake_st.session_state, "home_user_confirmation", True)
    monkeypatch.setitem(fake_st.session_state, "generated_prompt", "Test prompt")

    with patch("utils.runbook_generator_helpers.st", fake_st):
        render_prompt_preview(missing=[], section="home")

        fake_st.expander.assert_called_once_with("🧠 AI Prompt Preview (Optional)", expanded=True)
        fake_st.code.assert_called_once_with("Test prompt", language="markdown")
        fake_st.success.assert_called_once_with("✅ Prompt ready! Now you can generate your runbook.")

def test_render_prompt_preview_not_confirmed(monkeypatch):
    fake_st = MagicMock()
    fake_st.session_state = {
        "home_user_confirmation": False,
        "generated_prompt": None
    }

    with patch("utils.runbook_generator_helpers.st", fake_st):
        render_prompt_preview(missing=[], section="home")

        fake_st.info.assert_called_once_with("☕️ Please check the box to confirm AI prompt generation.")

def test_render_prompt_preview_missing_fields(monkeypatch):
    fake_st = MagicMock()
    fake_st.session_state = {
        "home_user_confirmation": True,
        "generated_prompt": None
    }

    with patch("utils.runbook_generator_helpers.st", fake_st):
        render_prompt_preview(missing=["Electricity Provider"], section="home")

        fake_st.warning.assert_called_once_with("⚠️ Cannot generate prompt. Missing: Electricity Provider")

def test_preview_runbook_output_with_content(monkeypatch):
    fake_st = MagicMock()
    fake_st.button.return_value = True  # Simulate button being clicked

    # Monkeypatch streamlit used in the module
    with patch("utils.runbook_generator_helpers.st", fake_st):
        runbook_text = "## Sample Runbook\n- Task A\n- Task B"
        preview_runbook_output(runbook_text)

        fake_st.button.assert_called_once_with("📖 Preview Runbook")
        fake_st.expander.assert_called_once_with("🧠 AI-Generated Runbook Preview", expanded=True)
        fake_st.markdown.assert_called_once_with(runbook_text)

def test_preview_runbook_output_no_content(monkeypatch):
    fake_st = MagicMock()

    with patch("utils.runbook_generator_helpers.st", fake_st):
        preview_runbook_output(runbook_text="")

        fake_st.warning.assert_called_once_with("⚠️ No runbook content available to preview.")
        fake_st.button.assert_not_called()

def test_maybe_render_download(monkeypatch):
    fake_st = MagicMock()
    fake_buffer = b"dummy docx content"
    fake_text = "## Sample Runbook"

    # Fake session state
    fake_st.session_state = {
        "runbook_buffer": fake_buffer,
        "runbook_text": fake_text
    }

    # Patch streamlit and preview function
    with patch("utils.runbook_generator_helpers.st", fake_st), \
         patch("utils.runbook_generator_helpers.preview_runbook_output") as mock_preview:
        
        maybe_render_download(section="home", filename=None)

        # Check that download button was rendered
        fake_st.download_button.assert_called_once_with(
            label="📥 Download DOCX",
            data=fake_buffer,
            file_name="home_emergency_runbook.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Check success message
        fake_st.success.assert_called_once_with("✅ Runbook ready for download!")

        # Check runbook preview
        mock_preview.assert_called_once_with(fake_text)

def test_maybe_render_download_no_buffer(monkeypatch):
    fake_st = MagicMock()
    fake_st.session_state = {
        "runbook_buffer": None,
        "runbook_text": "## Some Preview"
    }

    with patch("utils.runbook_generator_helpers.st", fake_st), \
         patch("utils.runbook_generator_helpers.preview_runbook_output") as mock_preview:

        maybe_render_download(section="pets", filename="custom.docx")

        # No download button if no buffer
        fake_st.download_button.assert_not_called()
        fake_st.success.assert_not_called()

        # Preview should still be rendered
        mock_preview.assert_called_once_with("## Some Preview")
